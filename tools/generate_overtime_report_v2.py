import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import time
import openpyxl
from openpyxl.styles import PatternFill
from docx import Document
from dotenv import load_dotenv
from google import genai

# ---------------------------------------------------------
# 初期設定 (Gemini API)
# ---------------------------------------------------------
env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
load_dotenv(env_path)
api_key = os.getenv("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)

# ---------------------------------------------------------
# 共通補助関数
# ---------------------------------------------------------
def load_csv_safely(filepath):
    for enc in ['cp932', 'utf-8', 'utf-8-sig', 'shift_jis']:
        try:
            return pd.read_csv(filepath, encoding=enc, engine='python', on_bad_lines='skip')
        except Exception:
            pass
    raise Exception(f"対応する文字コードで読み込めませんでした: {filepath}")

def load_master_safe(filepath):
    if filepath.lower().endswith('.csv'):
        return load_csv_safely(filepath)
    else:
        return pd.read_excel(filepath)

def get_audit_target_dates(excel_path, sheet_name='管理部提出用'):
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        if sheet_name not in wb.sheetnames: return set()
        ws = wb[sheet_name]
        
        name_col_idx = 2
        date_col_idx = 1
        
        for row in ws.iter_rows(min_row=1, max_row=5):
            for cell in row:
                val = str(cell.value).strip() if cell.value else ''
                if '端末' in val or '氏名' in val or '名前' in val: name_col_idx = cell.column
                if '日付' in val: date_col_idx = cell.column
            if name_col_idx and date_col_idx: break
                
        targets = set()
        ac_col_idx = 29  # AC列
        
        for row_idx in range(2, ws.max_row + 1):
            ac_cell = ws.cell(row=row_idx, column=ac_col_idx)
            val = str(ac_cell.value).strip() if ac_cell.value else ''
            
            if val in ['〇', '○', 'o', 'O', '丸', '対象']:
                name_val = str(ws.cell(row=row_idx, column=name_col_idx).value or '').strip()
                date_val = str(ws.cell(row=row_idx, column=date_col_idx).value or '').strip()
                if name_val and date_val:
                    try:
                        clean_date = date_val.split('（')[0].split('(')[0].strip()
                        dt = pd.to_datetime(clean_date).strftime('%m/%d')
                        name_clean = name_val.replace(' ', '').replace('　', '').replace('社員', '')
                        targets.add((name_clean, dt))
                    except: pass
        return targets
    except Exception as e:
        print(f"監査対象日（AC列の〇）の抽出に失敗しました: {e}")
        return set()

def format_est_time(diff_sec):
    if diff_sec <= 0: return "約0分"
    h, rem = divmod(int(diff_sec), 3600)
    m, _ = divmod(rem, 60)
    if h == 0: return f"約{m}分"
    elif m == 0: return f"約{h}時間"
    else: return f"約{h}時間{m:02d}分"

def find_col(df, keywords, fallback_index=None):
    for col in df.columns:
        for kw in keywords:
            if kw in col: return col
    if fallback_index is not None and fallback_index < len(df.columns): return df.columns[fallback_index]
    return df.columns[0]

def get_val(row, col_name, fallback_index=None):
    if col_name in row.index: return str(row[col_name])
    if fallback_index is not None and fallback_index < len(row.index): return str(row.iloc[fallback_index])
    return ''

def generate_summary_with_ai(log_df, col_title, col_path):
    if not client: return "APIキー未設定のため生成スキップ", "APIキー未設定"
    
    log_lines = []
    for _, row in log_df.iterrows():
        t = str(row.get(col_title, ''))
        p = str(row.get(col_path, ''))
        if t and t != 'nan' and '電源' not in t: log_lines.append(f"[{p}] {t}")
    
    unique_logs = list(dict.fromkeys(log_lines))[:100]
    log_text = "\n".join(unique_logs)
    if not log_text.strip(): return "具体的な作業ログなし", "特記事項なし"
        
    prompt = f"""
あなたは企業の監査担当です。以下のPC操作ログから、作業の詳細と総括サマリーを生成してください。
【操作ログ】
{log_text}

【出力フォーマット指示】
必ず以下の2つのセクションに分けて出力してください。不要な挨拶等は一切不要です。

■作業詳細
（Excelのファイル名やシステム名を元に、何を作業していたか箇条書きで簡潔に整理）

■サマリー
（上記を踏まえ、どんな業務を行っていたかを客観的で厳格な2〜3行の自然な日本語で総括してください）
"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            text = response.text
            details, summary = "詳細のパースに失敗しました。", "サマリーのパースに失敗しました。"
            if "■作業詳細" in text and "■サマリー" in text:
                parts = text.split("■サマリー")
                details = parts[0].replace("■作業詳細", "").strip()
                summary = parts[1].strip()
            else:
                summary = text.strip()
            return details, summary
        except Exception as e:
            if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                return "AI生成スキップ（API制限到達）", "API利用制限により要約を生成できませんでした。"
            elif '503' in str(e) or 'UNAVAILABLE' in str(e):
                if attempt < max_retries - 1: time.sleep(5)
                else: return "AI生成エラー: サーバー混雑", "現在AIサーバーが非常に混雑しています。"
            else:
                return f"AI生成エラー: {str(e)}", "エラーのため生成できませんでした"


# =========================================================
# 【STEP-1】: 管理部提出用シートの作成 (VBA代替機能)
# =========================================================
def step1_create_submission_list(csv_paths, master_path, output_dir):
    try:
        print("管理コンソール(台帳)を読み込んでいます...")
        df_master = load_master_safe(master_path)
        
        dict_terminal = {}
        dict_dept = {}
        for _, row in df_master.iterrows():
            if len(row) > 4:
                comp_name = str(row.iloc[4]).strip().upper() # E列: コンピューター名
                if comp_name and comp_name != 'NAN':
                    dict_terminal[comp_name] = str(row.iloc[3]) if len(row) > 3 else "" # D列: 端末機名(氏名)
                    dict_dept[comp_name] = str(row.iloc[9]) if len(row) > 9 else ""     # J列: 部署名
        
        print("SKYSEAログから電源ON/OFF時間を抽出中...")
        results = {}
        for csv_path in csv_paths:
            df_log = load_csv_safely(csv_path)
            for _, row in df_log.iterrows():
                if len(row) > 11:
                    dt_val = str(row.iloc[6]).strip()            # G列: 日時
                    comp_name = str(row.iloc[1]).strip().upper() # B列: コンピューター名
                    op_type = str(row.iloc[11])                  # L列: 操作種別
                    
                    if not dt_val or dt_val.lower() == 'nan' or not comp_name or comp_name.lower() == 'nan':
                        continue
                    
                    try:
                        dt = pd.to_datetime(dt_val)
                        # 翌朝5:00までのログは前日の日付として扱う
                        logical_dt = dt - pd.Timedelta(hours=5)
                        yobi = ["月", "火", "水", "木", "金", "土", "日"][logical_dt.weekday()]
                        date_str = logical_dt.strftime('%Y/%m/%d') + f"（{yobi}）"
                        key = (date_str, comp_name)
                        
                        if key not in results:
                            results[key] = [None, None]
                            
                        if "電源ON" in op_type:
                            if results[key][0] is None or dt < results[key][0]:
                                results[key][0] = dt
                        elif "電源OFF" in op_type:
                            if results[key][1] is None or dt > results[key][1]:
                                results[key][1] = dt
                    except:
                        pass
        
        print("出力データを構築中...")
        output_data = []
        for (date_str, comp_name), (time_on_dt, time_off_dt) in results.items():
            time_on = time_on_dt.strftime('%H:%M:%S') if time_on_dt is not None else ""
            time_off = time_off_dt.strftime('%H:%M:%S') if time_off_dt is not None else ""
            term_name = dict_terminal.get(comp_name, "(未登録)")
            dept_name = dict_dept.get(comp_name, "-")
            output_data.append([date_str, term_name, comp_name, dept_name, time_on, time_off])
            
        df_out = pd.DataFrame(output_data, columns=["日付", "端末機名", "コンピューター名", "部署名", "電源ON時間", "電源OFF時間"])
        # VBAの書き出しに近い形でソート（日付・コンピューター名など）
        df_out = df_out.sort_values(by=["日付", "部署名", "コンピューター名"])
        
        out_filename = f"管理部提出用リスト_{time.strftime('%Y%m%d%H%M%S')}.xlsx"
        out_filepath = os.path.join(output_dir, out_filename)
        
        df_out.to_excel(out_filepath, sheet_name="管理部提出用", index=False)
        
        # ヘッダーに色付け (RGB: 220, 230, 241 -> Hex: DCE6F1)
        wb = openpyxl.load_workbook(out_filepath)
        ws = wb["管理部提出用"]
        fill_color = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
        for cell in ws[1]:
            cell.fill = fill_color
        
        # 列幅の自動調整
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length: max_length = len(str(cell.value))
                except: pass
            ws.column_dimensions[column].width = (max_length + 2)

        wb.save(out_filepath)
        messagebox.showinfo("STEP-1 完了", f"VBAと同じ動作で【管理部提出用シート】を作成しました！\nまずはこのファイルを提出してください。\n\n出力先: {out_filepath}")
        
    except Exception as e:
        messagebox.showerror("エラー", f"処理中にエラーが発生しました:\n{e}")

# =========================================================
# 【STEP-2】: 残業調査報告書の自動生成 (従来の機能)
# =========================================================
def step2_generate_reports(excel_path, csv_paths, output_dir):
    print("提出用シート(勤怠マスタ)の読み込みを開始します（色の解析中...）")
    
    audit_targets = get_audit_target_dates(excel_path)
    if not audit_targets: print("[警告] AC列に「〇」（監査対象）が見つかりませんでした。全件処理します。")
    else: print(f"監査対象として {len(audit_targets)} 件の「〇」データを検出しました！")
    
    df_master = pd.read_excel(excel_path, sheet_name='管理部提出用')
    df_master.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_master.columns]
    
    # 端末機名(氏名)をキーとして探す
    name_col = find_col(df_master, ['端末', '氏名', '名前'], 1)
    
    print(f"\n合計 {len(csv_paths)} 個のログファイルを連続処理します。\n")

    for csv_path in csv_paths:
        filename_csv = os.path.basename(csv_path)
        print(f"========== 処理開始: {filename_csv} ==========")
        
        target_name_guess = None
        for name in df_master[name_col].dropna().unique():
            name_clean = str(name).replace(' ', '').replace('　', '').replace('社員', '')
            if name_clean and name_clean in filename_csv.replace(' ', '').replace('　', ''):
                target_name_guess = str(name)
                break
                
        if not target_name_guess:
            target_name_guess = filename_csv.split('_')[0] if '_' in filename_csv else filename_csv[:2]
            
        df_target_master = df_master[df_master[name_col].str.contains(target_name_guess, na=False)].copy()
        
        if len(df_target_master) == 0:
            print(f"  -> [スキップ] マスタから対象者「{target_name_guess}」が見つかりませんでした。")
            continue
            
        date_col = find_col(df_target_master, ['日付'], 0)
        clean_dates = df_target_master[date_col].astype(str).str.split('（').str[0].str.split('(').str[0].str.strip()
        df_target_master['日付_master'] = pd.to_datetime(clean_dates, errors='coerce')
        
        try: df_log = load_csv_safely(csv_path)
        except Exception as e:
            print(f"  -> [エラー] CSV読み込み失敗: {e}"); continue

        df_log.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_log.columns]
        col_datetime = find_col(df_log, ['日時'], 6)
        col_path = find_col(df_log, ['パス', 'URL', 'ＵＲＬ'], 11)
        col_title = df_log.columns[14] if len(df_log.columns) > 14 else find_col(df_log, ['タイトル'], 14)
        
        try:
            df_log['日時'] = pd.to_datetime(df_log[col_datetime])
            df_log = df_log.sort_values('日時').reset_index(drop=True)
            # 翌朝5:00までのログを前日扱いにする
            df_log['日付_log'] = (df_log['日時'] - pd.Timedelta(hours=5)).dt.normalize()
            
            df_log['次ログ日時'] = df_log['日時'].shift(-1)
            df_log['次ログ日付'] = df_log['日付_log'].shift(-1)
            df_log['操作秒数'] = (df_log['次ログ日時'] - df_log['日時']).dt.total_seconds().fillna(0)
            
            df_log.loc[df_log['日付_log'] != df_log['次ログ日付'], '操作秒数'] = 0
            df_log.loc[df_log['操作秒数'] < 0, '操作秒数'] = 0
            df_log['操作秒数'] = df_log['操作秒数'].apply(lambda x: min(x, 300))
        except Exception as e:
            print(f"  -> [エラー] 日時計算失敗: {e}"); continue

        daily_log = df_log.groupby('日付_log').agg(PC電源ON=('日時', 'min'), PC電源OFF=('日時', 'max')).reset_index()
        merged_df = pd.merge(df_target_master, daily_log, left_on='日付_master', right_on='日付_log', how='left')
        
        col_start = find_col(df_target_master, ['電源ON時間'], 4)
        col_end = find_col(df_target_master, ['電源OFF時間'], 5)

        processed_count = 0
        for index, row in merged_df.iterrows():
            if pd.isna(row['PC電源ON']): continue
                
            date_str = row['日付_master'].strftime('%Y/%m/%d')
            date_md = row['日付_master'].strftime('%m/%d')
            target_fullname = get_val(row, name_col)
            target_fullname_clean = target_fullname.replace(' ', '').replace('　', '').replace('社員', '')
            
            if audit_targets and (target_fullname_clean, date_md) not in audit_targets: continue
                
            processed_count += 1
            flags = []
            is_early_detected = False
            early_time_str = ""

            start_str = get_val(row, col_start)
            end_str = get_val(row, col_end)
            
            day_logs = df_log[df_log['日付_log'] == row['日付_log']]
            valid_logs = day_logs[day_logs[col_title].notna() & (day_logs[col_title].astype(str).str.strip() != '') & (day_logs[col_title].astype(str) != 'nan')]
            app_summary = valid_logs.groupby(col_title)['操作秒数'].sum().sort_values(ascending=False)
            
            print(f"    [{date_str}] Gemini APIでAIサマリーを生成中...")
            ai_details, ai_summary = generate_summary_with_ai(day_logs, col_title, col_path)
            
            doc = Document()
            doc.add_heading('残業調査報告書', 0)
            yobi = ["月", "火", "水", "木", "金", "土", "日"][row['日付_master'].weekday()]
            
            calc_start = row['PC電源ON']
            calc_end = row['PC電源OFF']
            diff_sec = (calc_end - calc_start).total_seconds()
            est_time_str = f"{calc_start.strftime('%H:%M')} ～ {calc_end.strftime('%H:%M')}（{format_est_time(diff_sec)}）"
            
            doc.add_paragraph(f"対象者：{target_fullname}")
            doc.add_paragraph(f"部署名：{get_val(row, '部署名', 3)}")
            doc.add_paragraph(f"{row['日付_master'].strftime('%Y/%m/%d')}（{yobi}）")
            doc.add_paragraph(f"PCの稼働：{row['PC電源ON'].strftime('%H:%M')} ～ {row['PC電源OFF'].strftime('%H:%M')}（最終）")
            doc.add_paragraph(f"作業していたと推測される時間：{est_time_str}")
            
            doc.add_heading('主な操作内容（タイトル別集計）', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = '操作タイトル'
            table.rows[0].cells[1].text = '合計時間（概算）'
            
            count = 0
            for app_name, total_sec in app_summary.items():
                if count >= 10 or total_sec == 0: break
                row_cells = table.add_row().cells
                title_text = str(app_name)
                if len(title_text) > 40: title_text = title_text[:38] + "..."
                row_cells[0].text = title_text
                
                h_app, rem_app = divmod(int(total_sec), 3600)
                m_app, s_app = divmod(rem_app, 60)
                
                if h_app == 0: row_cells[1].text = f"{m_app}分{s_app}秒"
                elif m_app == 0 and s_app == 0: row_cells[1].text = f"{h_app}時間"
                else: row_cells[1].text = f"{h_app}時間{m_app:02d}分"
                count += 1
                
            doc.add_heading('作業詳細', level=2)
            doc.add_paragraph(ai_details)
            doc.add_heading('サマリー', level=2)
            doc.add_paragraph(ai_summary)
            
            target_name_short = target_fullname.replace('社員', '').replace(' ', '').replace('　', '')
            safe_date = date_str.replace('/', '')
            filename = f"残業調査報告書_{target_name_short}_{safe_date}.docx"
            filepath = os.path.join(output_dir, filename)
            
            try:
                doc.save(filepath)
                print(f"      -> 保存完了: {filename}")
            except Exception as e: print(f"      -> [エラー] 保存失敗: {e}")
            time.sleep(1)

    print("\n全ての処理が完了しました！")
    messagebox.showinfo("STEP-2 完了", "すべてのWord報告書の生成が完了しました！")

# =========================================================
# GUI 統合アプリケーション
# =========================================================
class OvertimeSystemApp_V2:
    def __init__(self, root):
        self.root = root
        self.root.title("残業調査 オールインワンシステム Ver.2")
        self.root.geometry("600x300")
        self.root.attributes('-topmost', True)
        
        tk.Label(root, text="監査業務 オールインワン処理メニュー", font=("Meiryo", 14, "bold")).pack(pady=15)
        
        btn1 = tk.Button(root, text="STEP-1: 【管理部提出用シート】の自動作成 (VBA完全代替)", command=self.run_step1, width=55, height=2, bg="#e0f7fa", font=("Meiryo", 10))
        btn1.pack(pady=10)
        
        btn2 = tk.Button(root, text="STEP-2: 残業調査報告書(Word)の自動生成", command=self.run_step2, width=55, height=2, bg="#fce4ec", font=("Meiryo", 10))
        btn2.pack(pady=10)

    def run_step1(self):
        messagebox.showinfo("STEP-1", "【1/3】SKYSEAから抽出した「生ログ(CSV/Excel)」を選択してください。(複数可)")
        log_csv = filedialog.askopenfilenames(title="SKYSEA生ログ", filetypes=[("ログファイル", "*.csv *.xlsx *.xls")])
        if not log_csv: return
        
        messagebox.showinfo("STEP-1", "【2/3】「管理コンソールメイン画面(台帳)」のファイルを選択してください。")
        master_excel = filedialog.askopenfilename(title="管理コンソール台帳", filetypes=[("Excel/CSV", "*.xlsx *.xls *.csv")])
        if not master_excel: return
        
        messagebox.showinfo("STEP-1", "【3/3】完成した「管理部提出用リスト」を保存するフォルダを選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        step1_create_submission_list(log_csv, master_excel, out_dir)

    def run_step2(self):
        messagebox.showinfo("STEP-2", "【1/3】管理部から返却された「〇」印付きの提出用リスト(Excel)を選択してください。")
        master_excel = filedialog.askopenfilename(title="管理部返却済みリスト(Excel)", filetypes=[("Excel", "*.xlsx *.xls *.xlsm")])
        if not master_excel: return
        
        messagebox.showinfo("STEP-2", "【2/3】対象者の「PC操作ログ(CSV)」を選択してください。(複数可)")
        csvs = filedialog.askopenfilenames(title="PC操作ログ(CSV)", filetypes=[("CSV", "*.csv")])
        if not csvs: return
        
        messagebox.showinfo("STEP-2", "【3/3】報告書(Word)を保存する「出力先フォルダ」を選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        step2_generate_reports(master_excel, csvs, out_dir)

def main():
    root = tk.Tk()
    app = OvertimeSystemApp_V2(root)
    root.mainloop()

if __name__ == "__main__":
    main()
