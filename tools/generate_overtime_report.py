import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import time
import openpyxl
from docx import Document
from dotenv import load_dotenv
from google import genai

# ---------------------------------------------------------
# 初期設定 (Gemini API)
# ---------------------------------------------------------
env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
load_dotenv(env_path)
api_key = os.getenv("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)

# ---------------------------------------------------------
# 共通補助関数
# ---------------------------------------------------------
def get_audit_target_dates(excel_path, sheet_name='管理部提出用'):
    """ExcelのAC列(29)に「〇」が入力されている「監査対象行」の氏名と日付(MM/DD)を抽出する"""
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        if sheet_name not in wb.sheetnames:
            return set()
        ws = wb[sheet_name]
        
        name_col_idx = None
        date_col_idx = None
        
        for row in ws.iter_rows(min_row=1, max_row=5):
            for cell in row:
                val = str(cell.value).strip() if cell.value else ''
                if '氏名' in val or '名前' in val:
                    name_col_idx = cell.column
                if '日付' in val:
                    date_col_idx = cell.column
            if name_col_idx and date_col_idx:
                break
                
        if not name_col_idx or not date_col_idx:
            name_col_idx = 2
            date_col_idx = 3
            
        targets = set()
        ac_col_idx = 29  # AC列
        
        for row_idx in range(2, ws.max_row + 1):
            ac_cell = ws.cell(row=row_idx, column=ac_col_idx)
            val = str(ac_cell.value).strip() if ac_cell.value else ''
            
            if val in ['〇', '○', 'o', 'O', '丸', '対象']:
                name_val = str(ws.cell(row=row_idx, column=name_col_idx).value or '').strip()
                date_val = ws.cell(row=row_idx, column=date_col_idx).value
                if name_val and date_val:
                    try:
                        dt = pd.to_datetime(date_val).strftime('%m/%d')
                        name_clean = name_val.replace(' ', '').replace('　', '').replace('社員', '')
                        targets.add((name_clean, dt))
                    except:
                        pass
        return targets
    except Exception as e:
        print(f"監査対象日（AC列の〇）の抽出に失敗しました: {e}")
        return set()

def format_est_time(diff_sec):
    if diff_sec <= 0: return "約0分"
    h, rem = divmod(int(diff_sec), 3600)
    m, _ = divmod(rem, 60)
    if h == 0:
        return f"約{m}分"
    elif m == 0:
        return f"約{h}時間"
    else:
        return f"約{h}時間{m:02d}分"

def load_csv_safely(filepath):
    for enc in ['cp932', 'utf-8', 'utf-8-sig', 'shift_jis']:
        try:
            return pd.read_csv(filepath, encoding=enc, engine='python', on_bad_lines='skip')
        except Exception:
            pass
    raise Exception(f"対応する文字コードで読み込めませんでした: {filepath}")

def find_col(df, keywords, fallback_index=None):
    for col in df.columns:
        for kw in keywords:
            if kw in col:
                return col
    if fallback_index is not None and fallback_index < len(df.columns):
        return df.columns[fallback_index]
    return df.columns[0]

def get_val(row, col_name, fallback_index=None):
    if col_name in row.index: return str(row[col_name])
    if fallback_index is not None and fallback_index < len(row.index): return str(row.iloc[fallback_index])
    return ''

def generate_summary_with_ai(log_df, col_title, col_path):
    if not client:
        return "APIキー未設定のため生成スキップ", "APIキー未設定"
    
    log_lines = []
    for _, row in log_df.iterrows():
        t = str(row.get(col_title, ''))
        p = str(row.get(col_path, ''))
        if t and t != 'nan' and '電源' not in t:
            log_lines.append(f"[{p}] {t}")
    
    unique_logs = list(dict.fromkeys(log_lines))[:100]
    log_text = "\n".join(unique_logs)
    
    if not log_text.strip():
        return "具体的な作業ログなし", "特記事項なし"
        
    prompt = f"""
あなたは企業の監査担当です。以下のPC操作ログから、作業の詳細と総括サマリーを生成してください。
【操作ログ】
{log_text}

【出力フォーマット指示】
必ず以下の2つのセクションに分けて出力してください。不要な挨拶等は一切不要です。

■作業詳細
（Excelのファイル名やシステム名を元に、何を作業していたか箇条書きで簡潔に整理）

■サマリー
（上記を踏まえ、どんな業務を行っていたかを客観的で厳格な2〜3行の自然な日本語で総括してください）
"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            text = response.text
            
            details, summary = "詳細のパースに失敗しました。", "サマリーのパースに失敗しました。"
            if "■作業詳細" in text and "■サマリー" in text:
                parts = text.split("■サマリー")
                details = parts[0].replace("■作業詳細", "").strip()
                summary = parts[1].strip()
            else:
                summary = text.strip()
                
            return details, summary
            
        except Exception as e:
            if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                print(f"  -> [API制限] Gemini APIの無料枠制限（回数上限）に達しました。AI生成をスキップします。")
                return "AI生成スキップ（API制限到達）", "API利用制限により要約を生成できませんでした。"
            elif '503' in str(e) or 'UNAVAILABLE' in str(e):
                if attempt < max_retries - 1:
                    time.sleep(5)
                else:
                    return "AI生成エラー: サーバー混雑", "現在AIサーバーが非常に混雑しています。"
            else:
                return f"AI生成エラー: {str(e)}", "エラーのため生成できませんでした"


# ---------------------------------------------------------
# STEP-1: 氏名クレンジング＆分割
# ---------------------------------------------------------
def step1_clean_and_split(master_excel, raw_csv_paths, output_dir):
    try:
        print("勤怠マスタを読み込んで氏名リストを作成中...")
        df_master = pd.read_excel(master_excel, sheet_name='管理部提出用')
        df_master.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_master.columns]
        name_col = find_col(df_master, ['氏名', '名前'], 1)
        
        master_names = df_master[name_col].dropna().unique()
        # 照合用の辞書 {空白抜きの氏名: 元の氏名}
        clean_names = {str(n).replace(' ','').replace('　','').replace('社員',''): str(n) for n in master_names}
        
        processed_count = 0
        for raw_csv in raw_csv_paths:
            print(f"ログファイルを処理中: {os.path.basename(raw_csv)}")
            df_log = load_csv_safely(raw_csv)
            # ユーザー名が入っていそうな列を探す（第5列目あたりをフォールバックに）
            user_col = find_col(df_log, ['ユーザ', '名前', '氏名', 'ログイン'], 4)
            
            for clean_name, original_name in clean_names.items():
                if clean_name == 'nan' or not clean_name: continue
                
                # ユーザー列にマスタの氏名（スペース抜き）が含まれているか判定
                mask = df_log[user_col].astype(str).str.replace(' ','').str.replace('　','').str.contains(clean_name, na=False)
                person_df = df_log[mask]
                
                if len(person_df) > 0:
                    out_name = f"{clean_name}_クレンジング済ログ.csv"
                    out_path = os.path.join(output_dir, out_name)
                    person_df.to_csv(out_path, index=False, encoding='cp932')
                    processed_count += 1
                    print(f"  -> {out_name} を出力しました。 ({len(person_df)}行)")
                    
        messagebox.showinfo("STEP-1 完了", f"名寄せと分割が完了しました。\n合計 {processed_count} 名分のCSVを出力しました。")
    except Exception as e:
        messagebox.showerror("エラー", f"STEP-1 でエラーが発生しました:\n{e}")

# ---------------------------------------------------------
# STEP-2: PC稼働時間一覧表の作成
# ---------------------------------------------------------
def step2_generate_pc_uptime_list(csv_paths, output_dir):
    try:
        results = []
        print("PC稼働時間一覧表を作成中...")
        for csv_path in csv_paths:
            df = load_csv_safely(csv_path)
            col_datetime = find_col(df, ['日時'], 6)
            col_pc = find_col(df, ['コンピュータ', 'PC', 'ホスト'], 1)
            
            df['日時_tmp'] = pd.to_datetime(df[col_datetime], errors='coerce')
            df = df.dropna(subset=['日時_tmp'])
            df['日付'] = df['日時_tmp'].dt.date
            
            grouped = df.groupby('日付').agg(
                PC起動時間=('日時_tmp', 'min'),
                PC終了時間=('日時_tmp', 'max')
            ).reset_index()
            
            # ファイル名から氏名を推測（STEP1で "_クレンジング済ログ" と付いている前提）
            filename = os.path.basename(csv_path).replace('_クレンジング済ログ.csv', '').replace('.csv', '')
            pc_name = df[col_pc].iloc[0] if col_pc in df.columns and len(df) > 0 else '不明'
            
            for _, row in grouped.iterrows():
                results.append({
                    '氏名': filename,
                    'PC名': pc_name,
                    '日付': row['日付'],
                    'PC起動時間': row['PC起動時間'].strftime('%H:%M:%S'),
                    'PC終了時間': row['PC終了時間'].strftime('%H:%M:%S')
                })
                
        if results:
            out_excel = os.path.join(output_dir, "STEP2_PC稼働時間一覧表.xlsx")
            pd.DataFrame(results).to_excel(out_excel, index=False)
            print(f"-> 完了: {out_excel}")
            messagebox.showinfo("STEP-2 完了", f"PC稼働時間一覧表を作成しました。\n出力先: {out_excel}")
        else:
            messagebox.showwarning("警告", "有効なデータが見つかりませんでした。")
    except Exception as e:
        messagebox.showerror("エラー", f"STEP-2 でエラーが発生しました:\n{e}")

# ---------------------------------------------------------
# STEP-3: 残業調査報告書の生成 (従来のメイン処理)
# ---------------------------------------------------------
def step3_generate_reports(excel_path, csv_paths, output_dir):
    print("勤怠マスタの読み込みを開始します（色の解析中...）")
    
    audit_targets = get_audit_target_dates(excel_path)
    if not audit_targets:
        print("[警告] マスタのAC列に「〇」（監査対象）が見つかりませんでした。全件処理します。")
    else:
        print(f"監査対象として {len(audit_targets)} 件の「〇」データを検出しました！")
    
    df_master = pd.read_excel(excel_path, sheet_name='管理部提出用')
    df_master.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_master.columns]
    name_col = find_col(df_master, ['氏名', '名前'], 1)
    
    print(f"\n合計 {len(csv_paths)} 個のログファイルを連続処理します。\n")

    for csv_path in csv_paths:
        filename_csv = os.path.basename(csv_path)
        print(f"========== 処理開始: {filename_csv} ==========")
        
        target_name_guess = None
        for name in df_master[name_col].dropna().unique():
            name_clean = str(name).replace(' ', '').replace('　', '').replace('社員', '')
            if name_clean and name_clean in filename_csv.replace(' ', '').replace('　', ''):
                target_name_guess = str(name)
                break
                
        if not target_name_guess:
            target_name_guess = filename_csv.split('_')[0] if '_' in filename_csv else filename_csv[:2]
            
        df_target_master = df_master[df_master[name_col].str.contains(target_name_guess, na=False)].copy()
        
        if len(df_target_master) == 0:
            print(f"  -> [スキップ] 勤怠マスタから対象者「{target_name_guess}」が見つかりませんでした。")
            continue
            
        date_col = find_col(df_target_master, ['日付'], 2)
        df_target_master['日付_master'] = pd.to_datetime(df_target_master[date_col])
        
        try:
            df_log = load_csv_safely(csv_path)
        except Exception as e:
            print(f"  -> [エラー] CSV読み込み失敗: {e}")
            continue

        df_log.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_log.columns]
        
        col_datetime = find_col(df_log, ['日時'], 6)
        col_path = find_col(df_log, ['パス', 'URL', 'ＵＲＬ'], 11)
        
        if len(df_log.columns) > 14:
            col_title = df_log.columns[14]
        else:
            col_title = find_col(df_log, ['タイトル'], 14)
        
        try:
            df_log['日時'] = pd.to_datetime(df_log[col_datetime])
            df_log = df_log.sort_values('日時').reset_index(drop=True)
            df_log['日付_log'] = df_log['日時'].dt.normalize()
            
            df_log['次ログ日時'] = df_log['日時'].shift(-1)
            df_log['次ログ日付'] = df_log['日付_log'].shift(-1)
            df_log['操作秒数'] = (df_log['次ログ日時'] - df_log['日時']).dt.total_seconds().fillna(0)
            
            df_log.loc[df_log['日付_log'] != df_log['次ログ日付'], '操作秒数'] = 0
            df_log.loc[df_log['操作秒数'] < 0, '操作秒数'] = 0
            
            df_log['操作秒数'] = df_log['操作秒数'].apply(lambda x: min(x, 300))
            
        except Exception as e:
            print(f"  -> [エラー] 日時データの計算に失敗しました: {e}")
            continue

        daily_log = df_log.groupby('日付_log').agg(PC電源ON=('日時', 'min'), PC電源OFF=('日時', 'max')).reset_index()
        merged_df = pd.merge(df_target_master, daily_log, left_on='日付_master', right_on='日付_log', how='left')
        
        col_status = find_col(df_target_master, ['勤務日種別', '区分'], 2)
        col_jiyu = find_col(df_target_master, ['事由'], 11)
        col_start = find_col(df_target_master, ['打刻・出勤', '出勤'], 3)
        col_end = find_col(df_target_master, ['打刻・退勤', '退勤'], 4)

        processed_count = 0
        for index, row in merged_df.iterrows():
            if pd.isna(row['PC電源ON']): continue
                
            date_str = row['日付_master'].strftime('%Y/%m/%d')
            date_md = row['日付_master'].strftime('%m/%d')
            target_fullname = get_val(row, name_col)
            target_fullname_clean = target_fullname.replace(' ', '').replace('　', '').replace('社員', '')
            
            if audit_targets and (target_fullname_clean, date_md) not in audit_targets:
                continue
                
            processed_count += 1
            shift_info = f"{get_val(row, col_status)} {get_val(row, col_jiyu)}".strip().replace('nan', '')
            flags = []
            
            is_early_detected = False
            early_time_str = ""
            
            if ('休' in shift_info):
                flags.append(f"【休日稼働】シフトは「{shift_info}」ですがPCが稼働しています")
            if 'テレワーク' in shift_info:
                flags.append(f"【テレワーク検知】社外からの稼働です")

            start_str = get_val(row, col_start)
            end_str = get_val(row, col_end)
            
            if ('休' not in shift_info):
                if start_str and start_str != 'nan':
                    try:
                        work_start = pd.to_datetime(f"{date_str} {start_str}")
                        if row['PC電源ON'] < work_start:
                            diff = int((work_start - row['PC電源ON']).total_seconds() / 60)
                            if diff >= 10:
                                flags.append(f"【早出疑義】出勤打刻より {diff}分早く PC起動({row['PC電源ON'].strftime('%H:%M')})")
                                is_early_detected = True
                            diff_sec_e = (work_start - row['PC電源ON']).total_seconds()
                            early_time_str = f"{row['PC電源ON'].strftime('%H:%M')} ～ {work_start.strftime('%H:%M')}（{format_est_time(diff_sec_e)}）"
                    except: pass

                if end_str and end_str != 'nan':
                    try:
                        work_end = pd.to_datetime(f"{date_str} {end_str}")
                        if row['PC電源OFF'] > work_end + pd.Timedelta(minutes=10):
                            diff = int((row['PC電源OFF'] - work_end).total_seconds() / 60)
                            flags.append(f"【残業疑義】退勤打刻より {diff}分遅くまで PC稼働({row['PC電源OFF'].strftime('%H:%M')})")
                    except: pass

            day_logs = df_log[df_log['日付_log'] == row['日付_log']]
            valid_logs = day_logs[day_logs[col_title].notna() & (day_logs[col_title].astype(str).str.strip() != '') & (day_logs[col_title].astype(str) != 'nan')]
            app_summary = valid_logs.groupby(col_title)['操作秒数'].sum().sort_values(ascending=False)
            
            print(f"    [{date_str}] Gemini APIでAIサマリーを生成中...")
            ai_details, ai_summary = generate_summary_with_ai(day_logs, col_title, col_path)
            
            doc = Document()
            doc.add_heading('残業調査報告書', 0)
            yobi = ["月", "火", "水", "木", "金", "土", "日"][row['日付_master'].weekday()]
            
            is_holiday = ('休' in shift_info)
            if end_str and end_str != 'nan': end_time_display = end_str
            elif is_holiday: end_time_display = "休日"
            else: end_time_display = "記録なし"

            calc_start = row['PC電源ON']
            calc_end = row['PC電源OFF']
            
            if not is_holiday and end_str and end_str != 'nan':
                try:
                    work_end_dt = pd.to_datetime(f"{date_str} {end_str}")
                    if row['PC電源OFF'] > work_end_dt: calc_start = work_end_dt
                except: pass

            diff_sec = (calc_end - calc_start).total_seconds()
            est_time_str = f"{calc_start.strftime('%H:%M')} ～ {calc_end.strftime('%H:%M')}（{format_est_time(diff_sec)}）"
            
            doc.add_paragraph(f"対象者：{target_fullname}")
            doc.add_paragraph(f"{row['日付_master'].strftime('%Y/%m/%d')}（{yobi}）")
            
            if is_early_detected and start_str and start_str != 'nan':
                doc.add_paragraph(f"出勤打刻時間：{start_str}")
                
            doc.add_paragraph(f"退勤打刻時間：{end_time_display}")
            doc.add_paragraph(f"PCの稼働：{row['PC電源ON'].strftime('%H:%M')} ～ {row['PC電源OFF'].strftime('%H:%M')}（最終）")
            
            if is_early_detected:
                doc.add_paragraph(f"【出勤打刻前に作業していたと推測される時間】：{early_time_str}")
                doc.add_paragraph(f"【退勤打刻後に作業していたと推測される時間】：{est_time_str}")
            else:
                doc.add_paragraph(f"作業していたと推測される時間：{est_time_str}")
            
            doc.add_heading('判定結果', level=2)
            if not flags: doc.add_paragraph("・問題なし（適正な稼働）")
            else:
                for f in flags: doc.add_paragraph(f"・{f}")
                    
            doc.add_heading('主な操作内容（タイトル別集計）', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = '操作タイトル'
            table.rows[0].cells[1].text = '合計時間（概算）'
            
            count = 0
            for app_name, total_sec in app_summary.items():
                if count >= 10 or total_sec == 0: break
                row_cells = table.add_row().cells
                title_text = str(app_name)
                if len(title_text) > 40: title_text = title_text[:38] + "..."
                row_cells[0].text = title_text
                
                h_app, rem_app = divmod(int(total_sec), 3600)
                m_app, s_app = divmod(rem_app, 60)
                
                if h_app == 0: row_cells[1].text = f"{m_app}分{s_app}秒"
                elif m_app == 0 and s_app == 0: row_cells[1].text = f"{h_app}時間"
                else: row_cells[1].text = f"{h_app}時間{m_app:02d}分"
                count += 1
                
            doc.add_heading('作業詳細', level=2)
            doc.add_paragraph(ai_details)
            doc.add_heading('サマリー', level=2)
            doc.add_paragraph(ai_summary)
            
            target_name_short = target_fullname.replace('社員', '').replace(' ', '').replace('　', '')
            safe_date = date_str.replace('/', '')
            filename = f"残業調査報告書_{target_name_short}_{safe_date}.docx"
            filepath = os.path.join(output_dir, filename)
            
            try:
                doc.save(filepath)
                print(f"      -> 保存完了: {filename}")
            except Exception as e:
                print(f"      -> [エラー] 保存失敗: {e}")
            time.sleep(1)

    print("\n全ての処理が完了しました！")
    messagebox.showinfo("STEP-3 完了", "すべてのWord報告書の生成が完了しました！\n出力先フォルダをご確認ください。")


# ---------------------------------------------------------
# GUI 統合アプリケーションクラス
# ---------------------------------------------------------
class OvertimeSystemApp:
    def __init__(self, root):
        self.root = root
        self.root.title("残業調査 オールインワンシステム")
        self.root.geometry("550x350")
        
        # 常に最前面に表示
        self.root.attributes('-topmost', True)
        
        tk.Label(root, text="監査業務 オールインワン処理メニュー", font=("Meiryo", 14, "bold")).pack(pady=15)
        
        btn1 = tk.Button(root, text="STEP-1: 生ログから個人別CSVに分割 (名寄せクレンジング)", command=self.run_step1, width=50, height=2, bg="#e0f7fa", font=("Meiryo", 10))
        btn1.pack(pady=8)
        
        btn2 = tk.Button(root, text="STEP-2: 個人別 PC稼働時間一覧表(Excel)の作成", command=self.run_step2, width=50, height=2, bg="#fff9c4", font=("Meiryo", 10))
        btn2.pack(pady=8)
        
        btn3 = tk.Button(root, text="STEP-3: 残業調査報告書(Word)の自動生成", command=self.run_step3, width=50, height=2, bg="#fce4ec", font=("Meiryo", 10))
        btn3.pack(pady=8)

    def run_step1(self):
        messagebox.showinfo("STEP-1", "【1/3】全社員の「勤怠マスタ(Excel)」を選択してください。")
        master_excel = filedialog.askopenfilename(title="勤怠マスタ(Excel)", filetypes=[("Excel", "*.xlsx *.xls")])
        if not master_excel: return
        
        messagebox.showinfo("STEP-1", "【2/3】SkySea等から出した「生ログ(CSV)」を選択してください。(複数可)")
        raw_csvs = filedialog.askopenfilenames(title="生ログ(CSV)", filetypes=[("CSV", "*.csv")])
        if not raw_csvs: return
        
        messagebox.showinfo("STEP-1", "【3/3】分割した個人CSVを保存する「出力先フォルダ」を選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        step1_clean_and_split(master_excel, raw_csvs, out_dir)

    def run_step2(self):
        messagebox.showinfo("STEP-2", "【1/2】対象者の「PC操作ログ(CSV)」を選択してください。(複数可)\n※STEP-1で分割したファイルなどを指定します。")
        csvs = filedialog.askopenfilenames(title="PC操作ログ(CSV)", filetypes=[("CSV", "*.csv")])
        if not csvs: return
        
        messagebox.showinfo("STEP-2", "【2/2】一覧表(Excel)を保存する「出力先フォルダ」を選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        step2_generate_pc_uptime_list(csvs, out_dir)

    def run_step3(self):
        messagebox.showinfo("STEP-3", "【1/3】全社員の「勤怠マスタ(Excel)」を選択してください。")
        master_excel = filedialog.askopenfilename(title="勤怠マスタ(Excel)", filetypes=[("Excel", "*.xlsx *.xls")])
        if not master_excel: return
        
        messagebox.showinfo("STEP-3", "【2/3】対象者の「PC操作ログ(CSV)」を選択してください。(複数可)")
        csvs = filedialog.askopenfilenames(title="PC操作ログ(CSV)", filetypes=[("CSV", "*.csv")])
        if not csvs: return
        
        messagebox.showinfo("STEP-3", "【3/3】報告書(Word)を保存する「出力先フォルダ」を選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        step3_generate_reports(master_excel, csvs, out_dir)

def main():
    root = tk.Tk()
    app = OvertimeSystemApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
