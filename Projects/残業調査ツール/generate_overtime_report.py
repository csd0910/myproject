import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import time
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from google import genai

# ---------------------------------------------------------
# 初期設定 (Gemini API)
# ---------------------------------------------------------
env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
load_dotenv(env_path)
api_key = os.getenv("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)

# ---------------------------------------------------------
# GUI ダイアログ関数
# ---------------------------------------------------------
def select_paths():
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    messagebox.showinfo("ファイル選択 (1/3)", "全社員の「勤怠マスタ（Excel）」を1つ選択してください。")
    excel_path = filedialog.askopenfilename(title="勤怠マスタ(Excel)を選択", filetypes=[("Excelファイル", "*.xlsx *.xls")])
    if not excel_path: return None, None, None

    messagebox.showinfo("ファイル選択 (2/3)", "対象者の「PC操作ログ（CSV）」を【複数選択】してください。\n※ShiftキーやCtrlキーを押しながらクリックすると複数ファイルを選べます。")
    csv_paths = filedialog.askopenfilenames(title="PC操作ログ(CSV)を複数選択", filetypes=[("CSVファイル", "*.csv")])
    if not csv_paths: return None, None, None

    messagebox.showinfo("フォルダ選択 (3/3)", "生成された報告書（Word）を保存する「出力先フォルダ」を選択してください。")
    output_dir = filedialog.askdirectory(title="報告書の保存先フォルダを選択")
    if not output_dir: return None, None, None

    return excel_path, csv_paths, output_dir

# ---------------------------------------------------------
# 解析補助関数
# ---------------------------------------------------------
def get_audit_target_dates(excel_path, sheet_name='管理部提出用'):
    """ExcelのAC列(29)に「〇」が入力されている「監査対象行」の氏名と日付(MM/DD)を抽出する"""
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        if sheet_name not in wb.sheetnames:
            return set()
        ws = wb[sheet_name]
        
        name_col_idx = None
        date_col_idx = None
        
        for row in ws.iter_rows(min_row=1, max_row=5):
            for cell in row:
                val = str(cell.value).strip() if cell.value else ''
                if '氏名' in val or '名前' in val:
                    name_col_idx = cell.column
                if '日付' in val:
                    date_col_idx = cell.column
            if name_col_idx and date_col_idx:
                break
                
        if not name_col_idx or not date_col_idx:
            name_col_idx = 2
            date_col_idx = 3
            
        targets = set()
        ac_col_idx = 29  # AC列 (A=1, Z=26, AA=27, AB=28, AC=29)
        
        for row_idx in range(2, ws.max_row + 1):
            ac_cell = ws.cell(row=row_idx, column=ac_col_idx)
            val = str(ac_cell.value).strip() if ac_cell.value else ''
            
            if val in ['〇', '○', 'o', 'O', '丸', '対象']:
                name_val = str(ws.cell(row=row_idx, column=name_col_idx).value or '').strip()
                date_val = ws.cell(row=row_idx, column=date_col_idx).value
                if name_val and date_val:
                    try:
                        dt = pd.to_datetime(date_val).strftime('%m/%d')
                        name_clean = name_val.replace(' ', '').replace('　', '').replace('社員', '')
                        targets.add((name_clean, dt))
                    except:
                        pass
        return targets
    except Exception as e:
        print(f"監査対象日（AC列の〇）の抽出に失敗しました: {e}")
        return set()

def format_est_time(diff_sec):
    if diff_sec <= 0: return "約0分"
    h, rem = divmod(int(diff_sec), 3600)
    m, _ = divmod(rem, 60)
    if h == 0:
        return f"約{m}分"
    elif m == 0:
        return f"約{h}時間"
    else:
        return f"約{h}時間{m:02d}分"

def load_csv_safely(filepath):
    for enc in ['cp932', 'utf-8', 'utf-8-sig', 'shift_jis']:
        try:
            return pd.read_csv(filepath, encoding=enc, engine='python', on_bad_lines='skip')
        except Exception:
            pass
    raise Exception("対応する文字コード(cp932, utf-8等)で読み込めませんでした。")

def find_col(df, keywords, fallback_index=None):
    for col in df.columns:
        for kw in keywords:
            if kw in col:
                return col
    if fallback_index is not None and fallback_index < len(df.columns):
        return df.columns[fallback_index]
    return df.columns[0]

def get_val(row, col_name, fallback_index=None):
    if col_name in row.index: return str(row[col_name])
    if fallback_index is not None and fallback_index < len(row.index): return str(row.iloc[fallback_index])
    return ''

# ---------------------------------------------------------
# Gemini API 連携関数
# ---------------------------------------------------------
def generate_summary_with_ai(log_df, col_title, col_path, is_early=False, is_overtime=False, 
                            start_str="記録なし", end_str="記録なし", pc_on_str="記録なし", pc_off_str="記録なし", last_human_op_str="記録なし"):
    if not client:
        return "APIキー未設定のため生成スキップ", "APIキー未設定"
    
    log_lines = []
    for _, row in log_df.iterrows():
        t = str(row.get(col_title, ''))
        p = str(row.get(col_path, ''))
        if t and t != 'nan' and '電源' not in t:
            log_lines.append(f"[{p}] {t}")
    
    unique_logs = list(dict.fromkeys(log_lines))[:100]
    log_text = "\n".join(unique_logs)
    
    if not log_text.strip():
        return "具体的な作業ログなし", "特記事項なし"
        
    prompt = f"""
あなたは企業の監査担当です。以下のPC操作ログと基本情報から、調査報告書用の「作業詳細」と「サマリー」を生成してください。

【基本情報】
・出勤打刻時刻：{start_str}
・退勤打刻時刻：{end_str}
・PC電源ON時刻：{pc_on_str}
・PC電源OFF時刻：{pc_off_str}
・人間による実操作（タイトルありログ）の最終時刻：{last_human_op_str}
・15分以上の早出疑義があるか：{'はい' if is_early else 'いいえ'}
・15分以上の残業疑義があるか：{'はい' if is_overtime else 'いいえ'}

【操作ログ抜粋（時間外の可能性が高い時間帯を中心に抽出）】
{log_text}

【出力フォーマット指示】
必ず以下の2つのセクションに分けて出力してください。

■作業詳細
（時間外のログを中心に、どのようなアプリやファイルを使用していたか箇条書きで簡潔に整理してください）

■サマリー
以下の形式に沿って、客観的かつ厳格に記述してください。

（記述例1：早出・残業ありの場合）
「出勤の打刻時間は {start_str} となっているが、PCは {pc_on_str} に起動し、[具体的な作業内容]などの作業が確認されました。これにより、出勤時刻より[差分時間]早く時間外作業が行われています。また、退勤の打刻時間は {end_str} となっているが、PCは {pc_off_str} まで稼働しており、打刻後に時間外作業が行われている事実が確認されました。この時間帯には、主に[具体的な作業内容]が行われていました。」

（記述例2：15分以内のため認められない場合）
「退勤の打刻時間は {end_str} となっているが、PCは {pc_off_str} に電源OFFされており、実操作の最終時刻も {last_human_op_str} であることから、15分以内で作業を終了しており残業は認められません。」

※判定結果は、必ず「15分ルール」に基づいて、「早出/残業と認められるか、あるいは認められないか」を明確に結論づけてください。
"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            print(f"    [Gemini API] 生成リクエスト送信中 (gemini-2.5-flash) - 試行 {attempt + 1}/{max_retries}...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )
            full_text = response.text
            if "■サマリー" in full_text:
                parts = full_text.split("■サマリー")
                details = parts[0].replace("■作業詳細", "").strip()
                summary = parts[1].strip()
                return details, summary
            return "（詳細生成失敗）", full_text
            
        except Exception as e:
            err_msg = str(e).upper()
            # 503 (UNAVAILABLE) や 429 (RESOURCE_EXHAUSTED) の場合にリトライ
            if any(code in err_msg for code in ["503", "UNAVAILABLE", "429", "RESOURCE_EXHAUSTED", "QUOTA"]):
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 5
                    print(f"      -> [API混雑] サーバーが混雑しています。{wait_time}秒後に再試行します...")
                    time.sleep(wait_time)
                    continue
            return f"AI生成エラー: {e}", "要約の生成中にエラーが発生しました。時間を置いて再度お試しください。"

# ---------------------------------------------------------
# メイン処理（レポート一括生成）
# ---------------------------------------------------------
def generate_reports(excel_path, csv_paths, output_dir):
    print("勤怠マスタの読み込みを開始します（色の解析中...）")
    
    audit_targets = get_audit_target_dates(excel_path)
    if not audit_targets:
        print("[警告] マスタのAC列に「〇」（監査対象）が見つかりませんでした。全件処理します。")
    else:
        print(f"監査対象として {len(audit_targets)} 件の「〇」データを検出しました！ (例: {list(audit_targets)[:3]})")
    
    df_master = pd.read_excel(excel_path, sheet_name='管理部提出用')
    df_master.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_master.columns]
    name_col = find_col(df_master, ['氏名', '名前'], 1)
    
    print(f"\n合計 {len(csv_paths)} 個のログファイルを連続処理します。\n")

    for csv_path in csv_paths:
        filename_csv = os.path.basename(csv_path)
        print(f"========== 処理開始: {filename_csv} ==========")
        
        target_name_guess = None
        for name in df_master[name_col].dropna().unique():
            name_clean = str(name).replace(' ', '').replace('　', '').replace('社員', '')
            if name_clean and name_clean in filename_csv.replace(' ', '').replace('　', ''):
                target_name_guess = str(name)
                break
                
        if not target_name_guess:
            target_name_guess = filename_csv.split('_')[0] if '_' in filename_csv else filename_csv[:2]
            
        df_target_master = df_master[df_master[name_col].str.contains(target_name_guess, na=False)].copy()
        
        if len(df_target_master) == 0:
            print(f"  -> [スキップ] 勤怠マスタから対象者「{target_name_guess}」が見つかりませんでした。")
            continue
            
        date_col = find_col(df_target_master, ['日付'], 2)
        df_target_master['日付_master'] = pd.to_datetime(df_target_master[date_col])
        
        try:
            df_log = load_csv_safely(csv_path)
        except Exception as e:
            print(f"  -> [エラー] CSV読み込み失敗: {e}")
            continue

        # ヘッダー名に含まれるゴミを掃除
        df_log.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_log.columns]
        
        col_datetime = find_col(df_log, ['日時'], 6)
        col_path = find_col(df_log, ['パス', 'URL', 'ＵＲＬ'], 11)
        
        # O列（14列目）をタイトルとして確実にとる
        if len(df_log.columns) > 14:
            col_title = df_log.columns[14]
        else:
            col_title = find_col(df_log, ['タイトル'], 14)
        
        try:
            df_log['日時'] = pd.to_datetime(df_log[col_datetime])
            df_log = df_log.sort_values('日時').reset_index(drop=True)
            df_log['日付_log'] = df_log['日時'].dt.normalize()
            
            # --- ユーザーの要望: G列(日時)の差分から操作時間を自力で計算する ---
            df_log['次ログ日時'] = df_log['日時'].shift(-1)
            df_log['次ログ日付'] = df_log['日付_log'].shift(-1)
            df_log['操作秒数'] = (df_log['次ログ日時'] - df_log['日時']).dt.total_seconds().fillna(0)
            
            # 日またぎ、または負の値の場合は0秒とする
            df_log.loc[df_log['日付_log'] != df_log['次ログ日付'], '操作秒数'] = 0
            df_log.loc[df_log['操作秒数'] < 0, '操作秒数'] = 0
            
            # 操作と操作の間が空きすぎている場合（例: 5分以上）は、PC放置とみなして5分(300秒)で打ち切る
            df_log['操作秒数'] = df_log['操作秒数'].apply(lambda x: min(x, 300))
            
        except Exception as e:
            print(f"  -> [エラー] 日時データの計算に失敗しました: {e}")
            continue

        daily_log = df_log.groupby('日付_log').agg(PC電源ON=('日時', 'min'), PC電源OFF=('日時', 'max')).reset_index()
        merged_df = pd.merge(df_target_master, daily_log, left_on='日付_master', right_on='日付_log', how='left')
        
        col_status = find_col(df_target_master, ['勤務日種別', '区分'], 2)
        col_jiyu = find_col(df_target_master, ['事由'], 11)
        # K列(10)を出勤、L列(11)を退勤として優先的に取得
        col_start = find_col(df_target_master, ['打刻・出勤', '出勤'], 10)
        col_end = find_col(df_target_master, ['打刻・退勤', '退勤'], 11)

        processed_count = 0
        for index, row in merged_df.iterrows():
            if pd.isna(row['PC電源ON']):
                continue
                
            date_str = row['日付_master'].strftime('%Y/%m/%d')
            date_md = row['日付_master'].strftime('%m/%d')
            target_fullname = get_val(row, name_col)
            target_fullname_clean = target_fullname.replace(' ', '').replace('　', '').replace('社員', '')
            
            if audit_targets and (target_fullname_clean, date_md) not in audit_targets:
                print(f"    -> [スキップ] {date_str} はマスタのAC列に〇がないためスキップしました。")
                continue
                
            processed_count += 1
            shift_info = f"{get_val(row, col_status)} {get_val(row, col_jiyu)}".strip().replace('nan', '')
            flags = []
            
            work_start = pd.NaT
            work_end = pd.NaT
            is_early_detected = False
            is_overtime_detected = False
            early_time_str = ""
            
            if ('休' in shift_info):
                flags.append(f"【休日稼働】シフトは「{shift_info}」ですがPCが稼働しています")
            
            if 'テレワーク' in shift_info:
                flags.append(f"【テレワーク検知】社外からの稼働です")

            start_str = get_val(row, col_start)
            end_str = get_val(row, col_end)
            
            if ('休' not in shift_info):
                if start_str and start_str != 'nan':
                    try:
                        work_start = pd.to_datetime(f"{date_str} {start_str}")
                        if row['PC電源ON'] < work_start:
                            diff = int((work_start - row['PC電源ON']).total_seconds() / 60)
                            if diff >= 15: # 15分以上前なら早出
                                flags.append(f"【早出疑義】出勤打刻の {diff}分前 からPC稼働({row['PC電源ON'].strftime('%H:%M')})が確認されました。")
                                is_early_detected = True
                            
                            diff_sec_e = (work_start - row['PC電源ON']).total_seconds()
                            early_time_str = f"{row['PC電源ON'].strftime('%H:%M')} ～ {work_start.strftime('%H:%M')}（約{format_est_time(diff_sec_e)}）"
                    except: pass

                if end_str and end_str != 'nan':
                    try:
                        work_end = pd.to_datetime(f"{date_str} {end_str}")
                        if pd.notna(work_end):
                            # 翌朝（最大12時間後）まで連続したログがあるか確認
                            actual_pc_off = row['PC電源OFF']
                            if actual_pc_off > work_end + pd.Timedelta(minutes=15):
                                diff = int((actual_pc_off - work_end).total_seconds() / 60)
                                flags.append(f"【残業疑義】退勤打刻の {diff}分後 までPC稼働({actual_pc_off.strftime('%H:%M')})が確認されました。")
                                is_overtime_detected = True
                                diff_sec_o = (actual_pc_off - work_end).total_seconds()
                                est_time_str = f"{work_end.strftime('%H:%M')} ～ {actual_pc_off.strftime('%H:%M')}（約{format_est_time(diff_sec_o)}）"
                            else:
                                is_overtime_detected = False
                    except: pass

            day_logs = df_log[df_log['日付_log'] == row['日付_log']]
            
            # --- ユーザー要望: O列（タイトル）がある項目だけをまとめる ---
            valid_logs = day_logs[day_logs[col_title].notna() & (day_logs[col_title].astype(str).str.strip() != '') & (day_logs[col_title].astype(str) != 'nan')]
            app_summary = valid_logs.groupby(col_title)['操作秒数'].sum().sort_values(ascending=False) if '操作秒数' in valid_logs.columns else {}

            # --- 実操作の最終時刻を特定し、更新処理を判定 ---
            if not valid_logs.empty:
                last_human_op_dt = valid_logs['日時'].max()
                last_human_op_str = last_human_op_dt.strftime('%H:%M')
                
                final_pc_dt = day_logs['日時'].max()
                # 実操作終了からPC終了まで10分以上の開きがある場合
                if (final_pc_dt - last_human_op_dt).total_seconds() > 600:
                    flags.append(f"【システム更新判定】実操作（タイトルありログ）は {last_human_op_str} までとなっております。")
                    flags.append(f"　それ以降のPC起動（{final_pc_dt.strftime('%H:%M')}まで）は、操作はしていないが更新処理などでPCを起動したままにしていたようだと推測され、実業務は認められませんでした。")
            else:
                if not day_logs.empty:
                    flags.append("【システム稼働判定】PCの起動は確認されましたが、人間による実操作（ログタイトルあり）は確認されませんでした。更新処理等の可能性があります。")
            
            # --- 打刻表示用テキストの準備 (AIサマリーでも使用) ---
            yobi = ["月", "火", "水", "木", "金", "土", "日"][row['日付_master'].weekday()]
            is_holiday = ('休' in shift_info)
            start_time_display = start_str if start_str and start_str != 'nan' else "記録なし"
            if end_str and end_str != 'nan':
                end_time_display = end_str
            elif is_holiday:
                end_time_display = "休日"
            else:
                end_time_display = "記録なし"

            print(f"    [{date_str}] Gemini APIでAIサマリーを生成中... (監査対象)")
            pc_on_str = row['PC電源ON'].strftime('%H:%M:%S')
            pc_off_str = row['PC電源OFF'].strftime('%H:%M:%S')
            ai_details, ai_summary = generate_summary_with_ai(
                day_logs, col_title, col_path, 
                is_early_detected, is_overtime_detected,
                start_time_display, end_time_display,
                pc_on_str, pc_off_str, last_human_op_str if 'last_human_op_str' in locals() else "記録なし"
            )
            
            doc = Document()
            title_p = doc.add_heading('残業調査報告書', 0)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            calc_start = row['PC電源ON']
            calc_end = row['PC電源OFF']
            
            if not is_holiday and end_str and end_str != 'nan':
                try:
                    work_end_dt = pd.to_datetime(f"{date_str} {end_str}")
                    if row['PC電源OFF'] > work_end_dt:
                        calc_start = work_end_dt
                except: pass

            diff_sec = (calc_end - calc_start).total_seconds()
            est_time_str = f"{calc_start.strftime('%H:%M')} ～ {calc_end.strftime('%H:%M')}（{format_est_time(diff_sec)}）"
            
            doc.add_paragraph(f"対象者：{target_fullname}")
            doc.add_paragraph(f"{row['日付_master'].strftime('%Y/%m/%d')}（{yobi}）")
            
            # 打刻情報の表示
            doc.add_paragraph(f"出勤打刻時間：{start_time_display}")
            doc.add_paragraph(f"退勤打刻時間：{end_time_display}")
            doc.add_paragraph(f"PCの稼働：{row['PC電源ON'].strftime('%H:%M')} ～ {row['PC電源OFF'].strftime('%H:%M')}（最終）")
            
            if is_early_detected:
                doc.add_paragraph(f"【出勤前の時間外と推測される時間】：{early_time_str}")
            else:
                doc.add_paragraph(f"【出勤前の時間外と推測される時間】：なし")

            if is_overtime_detected:
                doc.add_paragraph(f"【退勤後の時間外と推測される時間】：{est_time_str}")
            else:
                doc.add_paragraph(f"【退勤後の時間外と推測される時間】：なし")
            
            doc.add_heading('判定結果', level=2)
            if not flags: doc.add_paragraph("・問題なし（適正な稼働）")
            else:
                for f in flags: doc.add_paragraph(f"・{f}")
                    
            doc.add_heading('主な操作内容（タイトル別集計）', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = '操作タイトル'
            table.rows[0].cells[1].text = '合計時間（概算）'
            
            count = 0
            for app_name, total_sec in app_summary.items():
                if count >= 10 or total_sec == 0: break
                row_cells = table.add_row().cells
                
                # タイトルが長すぎる場合は切り詰め
                title_text = str(app_name)
                if len(title_text) > 40:
                    title_text = title_text[:38] + "..."
                row_cells[0].text = title_text
                
                h_app, rem_app = divmod(int(total_sec), 3600)
                m_app, s_app = divmod(rem_app, 60)
                
                if h_app == 0:
                    row_cells[1].text = f"{m_app}分{s_app}秒"
                elif m_app == 0 and s_app == 0:
                    row_cells[1].text = f"{h_app}時間"
                else:
                    row_cells[1].text = f"{h_app}時間{m_app:02d}分"
                    
                count += 1
                
            doc.add_heading('作業詳細', level=2)
            doc.add_paragraph(ai_details)
            doc.add_heading('サマリー', level=2)
            doc.add_paragraph(ai_summary)
            
            target_name_short = target_fullname.replace('社員', '').replace(' ', '').replace('　', '')
            safe_date = date_str.replace('/', '')
            filename = f"残業調査報告書_{target_name_short}_{safe_date}.docx"
            filepath = os.path.join(output_dir, filename)
            
            try:
                doc.save(filepath)
                print(f"      -> 保存完了: {filename}")
            except PermissionError:
                alt_filename = f"残業調査報告書_{target_name_short}_{safe_date}_再生成.docx"
                alt_filepath = os.path.join(output_dir, alt_filename)
                try:
                    doc.save(alt_filepath)
                    print(f"      -> [警告] ファイルが使用中のため、別名で保存しました: {alt_filename}")
                except Exception as e:
                    print(f"      -> [エラー] 保存に失敗しました: {e}")
            
            time.sleep(1)

        if processed_count == 0:
            print(f"    -> [スキップ] このファイルのログには、マスタのAC列に〇がついている監査対象日が含まれていませんでした。")

    print("\n全てのログファイルの処理が完了しました！出力先フォルダをご確認ください。")

def main():
    excel_path, csv_paths, output_dir = select_paths()
    if excel_path and csv_paths and output_dir:
        generate_reports(excel_path, csv_paths, output_dir)
    else:
        print("処理がキャンセルされました。")

if __name__ == "__main__":
    main()
