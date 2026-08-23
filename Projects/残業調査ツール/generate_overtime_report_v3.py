import pandas as pd
import tkinter as tk
import re
import tkinter.ttk as ttk
from tkinter import filedialog, messagebox
import os
import time
import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from dotenv import load_dotenv
from google import genai
import threading
import queue
import traceback

# ---------------------------------------------------------
# 定数・初期設定定義
# ---------------------------------------------------------
env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
load_dotenv(env_path)
api_key = os.getenv("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)

# Freee打刻検出用のURL/タイトル判定キーワード
FREEE_LOGIN_URL = "https://accounts.secure.freee.co.jp/sessions/new"
FREEE_LOGIN_TITLE = "freeeアカウントでログイン"
FREEE_STAMP_URL = "https://p.secure.freee.co.jp/#"
FREEE_STAMP_TITLE = "freee人事労務"

# ---------------------------------------------------------
# 共通補助関数
# ---------------------------------------------------------
def load_csv_safely(filepath):
    for enc in ['cp932', 'utf-8', 'utf-8-sig', 'shift_jis']:
        try:
            return pd.read_csv(filepath, encoding=enc, engine='python', on_bad_lines='skip')
        except Exception:
            pass
    raise Exception(f"対応する文字コードで読み込めませんでした: {filepath}")

def load_master_safe(filepath):
    if filepath.lower().endswith('.csv'):
        return load_csv_safely(filepath)
    else:
        return pd.read_excel(filepath)

def find_col(df, keywords, fallback_index=None):
    # まず完全一致で探す
    for col in df.columns:
        for kw in keywords:
            if str(col).strip() == str(kw).strip(): return col
    # 次に部分一致で探す
    for col in df.columns:
        for kw in keywords:
            if str(kw).strip() in str(col).strip(): return col
    if fallback_index is not None and fallback_index < len(df.columns): return df.columns[fallback_index]
    return df.columns[0]

def get_val(row, col_name, fallback_index=None):
    if col_name in row.index: return str(row[col_name])
    if fallback_index is not None and fallback_index < len(row.index): return str(row.iloc[fallback_index])
    return ''

def format_est_time(diff_sec):
    if diff_sec <= 0: return "約0分"
    h, rem = divmod(int(diff_sec), 3600)
    m, _ = divmod(rem, 60)
    if h == 0: return f"約{m}分"
    elif m == 0: return f"約{h}時間"
    else: return f"約{h}時間{m:02d}分"

def format_minutes_to_hm(minutes):
    if pd.isna(minutes) or minutes == "" or minutes is None:
        return ""
    try:
        mins = int(float(minutes))
        if mins <= 0:
            return ""
        h, m = divmod(mins, 60)
        if h == 0:
            return f"{m}分"
        else:
            return f"{h}時間{m}分"
    except:
        return ""

def normalize_name(name):
    if pd.isna(name):
        return ""
    name = str(name).strip()
    
    # 1. カッコ内の「前：」「前:」「元：」「元:」に続く名前を救出
    m = re.search(r'[\(（](?:前|元)[：:](.+?)[\)）]', name)
    if m:
        name = m.group(1)
    else:
        # 普通のカッコ書き（例：「会沢忍（OAセグメント）」）はカッコを除去
        name = re.sub(r'[\(（].*?[\)）]', '', name)
        
    # 2. 空白文字をすべて除去
    name = name.replace(' ', '').replace('　', '')
    name = re.sub(r'\s+', '', name)
    
    # 3. 異体字の標準化（表記揺れの置換）
    itaiji_map = {
        '邊': '辺', '邉': '辺',
        '齋': '斉', '齊': '斉', '斎': '斉',
        '髙': '高',
        '﨑': '崎',
        '嶋': '島',
        '栁': '柳',
        '澤': '沢',
    }
    for k, v in itaiji_map.items():
        name = name.replace(k, v)
        
    return name

def parse_time_with_crossover(date_str, time_str, base_time_dt=None):
    """
    日付と時間文字列をパースし、ベース時間より小さく日またぎが想定される場合は+1日する。
    """
    if pd.isna(time_str) or not str(time_str).strip() or str(time_str).lower() == 'nan':
        return pd.NaT
    
    t_str = str(time_str).strip()
    if len(t_str.split(':')) == 2:
        t_str += ":00"
        
    try:
        dt = pd.to_datetime(f"{date_str} {t_str}")
        if base_time_dt is not None and dt < base_time_dt:
            # 基準時間より前なら日付跨ぎとみなし、1日加算する
            dt = dt + pd.Timedelta(days=1)
        return dt
    except:
        return pd.NaT

def generate_summary_with_ai(log_df, col_title, col_path):
    if not client: return "APIキー未設定のため生成スキップ", "APIキー未設定"
    
    log_lines = []
    for _, row in log_df.iterrows():
        t = str(row.get(col_title, ''))
        p = str(row.get(col_path, ''))
        if t and t != 'nan' and '電源' not in t: log_lines.append(f"[{p}] {t}")
    
    unique_logs = list(dict.fromkeys(log_lines))[:100]
    log_text = "\n".join(unique_logs)
    if not log_text.strip(): return "具体的な作業ログなし", "特記事項なし"
        
    prompt = f"""
あなたは企業の監査担当です。以下のPC操作ログから、作業の詳細と総括サマリーを生成してください。
【操作ログ】
{log_text}

【出力フォーマット指示】
必ず以下の2つのセクションに分けて出力してください。不要な挨拶等は一切不要です。

■作業詳細
（Excelのファイル名やシステム名を元に、何を作業していたか箇条書きで簡潔に整理）

■サマリー
（上記を踏まえ、どんな業務を行っていたかを客観的で厳格な2〜3行の自然な日本語で総括してください）
"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            text = response.text
            details, summary = "詳細のパースに失敗しました。", "サマリーのパースに失敗しました。"
            if "■作業詳細" in text and "■サマリー" in text:
                parts = text.split("■サマリー")
                details = parts[0].replace("■作業詳細", "").strip()
                summary = parts[1].strip()
            else:
                summary = text.strip()
            return details, summary
        except Exception as e:
            if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                return "AI生成スキップ（API制限到達）", "API利用制限により要約を生成できませんでした。"
            elif '503' in str(e) or 'UNAVAILABLE' in str(e):
                if attempt < max_retries - 1: time.sleep(5)
                else: return "AI生成エラー: サーバー混雑", "現在AIサーバーが非常に混雑しています。"
            else:
                return f"AI生成エラー: {str(e)}", "エラーのため生成できませんでした"

# ---------------------------------------------------------
# 除外マスタのロードと自動生成機能
# ---------------------------------------------------------
def load_exclude_list(master_path, log_cb):
    """
    台帳と同じフォルダにある「除外マスタ.csv」をロードする。
    存在しない場合は、特定した共用端末や業務用端末のデフォルト除外リストで新規作成する。
    """
    master_dir = os.path.dirname(master_path)
    exclude_path = os.path.join(master_dir, "除外マスタ.csv")
    
    default_excludes = [
        # 第1弾 抽出分 (16件)
        {"除外対象": "FMV-MEET", "区分": "コンピューター名", "メモ": "福岡ミーティング卓"},
        {"除外対象": "WIN11-CTI", "区分": "コンピューター名", "メモ": "外部メディア接続端末"},
        {"除外対象": "WIN10-1135", "区分": "コンピューター名", "メモ": "物流共用"},
        {"除外対象": "FWN-032", "区分": "コンピューター名", "メモ": "物流部測定用（オフライン使用）"},
        {"除外対象": "FWN-028", "区分": "コンピューター名", "メモ": "名刺運用端末(福岡設置)"},
        {"除外対象": "DRVIEWER", "区分": "コンピューター名", "メモ": "DRviewer（オフライン使用）"},
        {"除外対象": "WIN10-1137", "区分": "コンピューター名", "メモ": "2Fミーティングスペース設置端末"},
        {"除外対象": "MS11", "区分": "コンピューター名", "メモ": "サーバ/マスターサーバー"},
        {"除外対象": "FWN-RELAYPC", "区分": "コンピューター名", "メモ": "SSL-VPN踏み台端末（大宮）"},
        {"除外対象": "DESKTOP-NDATC5R", "区分": "コンピューター名", "メモ": "SKY検証環境"},
        {"除外対象": "TIMEMGR1", "区分": "コンピューター名", "メモ": "タイムレコーダー共用端末TIMEMGR1"},
        {"除外対象": "TIMEMGR2", "区分": "コンピューター名", "メモ": "タイムレコーダー共用端末TIMEMGR2"},
        {"除外対象": "WIN10-CTI-TEST", "区分": "コンピューター名", "メモ": "外部メディア接続端末/テスト用"},
        {"除外対象": "F00133DL", "区分": "コンピューター名", "メモ": "福岡応接室端末"},
        {"除外対象": "WIN7-0053", "区分": "コンピューター名", "メモ": "物流部共用端末"},
        {"除外対象": "WIN10-1087", "区分": "コンピューター名", "メモ": "(川口)商品部共用"},
        # 第2弾 追加分 (7件)
        {"除外対象": "WIN10-1130", "区分": "コンピューター名", "メモ": "3F倉庫内端末(USB)"},
        {"除外対象": "WIN10-1040", "区分": "コンピューター名", "メモ": "FOREST-EDI端末"},
        {"除外対象": "WIN10-1010", "区分": "コンピューター名", "メモ": "銀行端末"},
        {"除外対象": "WIN10-1077", "区分": "コンピューター名", "メモ": "システム統括部共用"},
        {"除外対象": "FORESTOFFICE-PC1", "区分": "コンピューター名", "メモ": "貸出用ノート"},
        {"除外対象": "WIN10-1140", "区分": "コンピューター名", "メモ": "管理部・データ移動用"},
        {"除外対象": "WIN10-1038", "区分": "コンピューター名", "メモ": "FNETS障害調査用PC"},
    ]
    
    if not os.path.exists(exclude_path):
        try:
            log_cb(f"除外マスタが存在しないため新規作成します: {exclude_path}")
            df_default = pd.DataFrame(default_excludes)
            df_default.to_csv(exclude_path, index=False, encoding='cp932')
        except Exception as e:
            log_cb(f"[警告] 除外マスタの新規作成に失敗しました: {e}")
            
    excludes = set()
    if os.path.exists(exclude_path):
        try:
            df_ex = pd.read_csv(exclude_path, encoding='cp932')
            df_ex.columns = [str(c).strip().replace(' ', '') for c in df_ex.columns]
            for _, row in df_ex.iterrows():
                val = str(row.get("除外対象", "")).strip().upper()
                if val and val != "NAN":
                    excludes.add(val)
            log_cb(f"除外マスタから {len(excludes)} 件の除外キーを登録しました。")
        except Exception as e:
            log_cb(f"[エラー] 除外マスタの読み込みに失敗しました (フォールバック適用): {e}")
            
    if not excludes:
        excludes = set([d["除外対象"].upper() for d in default_excludes])
        log_cb(f"デフォルトの除外リスト {len(excludes)} 件を適用しました。")
        
    return excludes

def load_name_mapping(base_file_path, log_cb):
    """
    データファイルと同じフォルダにある「名前変換マスタ.csv」をロードする。
    存在しない場合は、同音異義字のサンプルをプリセットしたファイルを新規作成する。
    """
    file_dir = os.path.dirname(base_file_path)
    mapping_path = os.path.join(file_dir, "名前変換マスタ.csv")
    
    default_mappings = [
        {"マスタ名": "岡部正樹", "Freee名": "岡部政樹", "メモ": "同音異義字の表記揺れ補正"},
        {"マスタ名": "荒木一昌", "Freee名": "荒木一晶", "メモ": "同音異義字の表記揺れ補正"},
        {"マスタ名": "サンプル太郎", "Freee名": "サンプル 太郎", "メモ": "姓名の間にスペースがある場合などは自動除去されますが、個別指定も可能です"},
    ]
    
    if not os.path.exists(mapping_path):
        try:
            log_cb(f"名前変換マスタが存在しないため新規作成します: {mapping_path}")
            df_default = pd.DataFrame(default_mappings)
            df_default.to_csv(mapping_path, index=False, encoding='cp932')
        except Exception as e:
            log_cb(f"[警告] 名前変換マスタの新規作成に失敗しました: {e}")
            
    mapping_dict = {}
    if os.path.exists(mapping_path):
        try:
            df_map = None
            for enc in ['cp932', 'utf-8', 'utf-8-sig', 'shift_jis']:
                try:
                    df_map = pd.read_csv(mapping_path, encoding=enc, engine='python', on_bad_lines='skip')
                    break
                except:
                    pass
            
            if df_map is not None:
                df_map.columns = [str(c).strip().replace(' ', '').replace('　', '') for c in df_map.columns]
                col_master = find_col(df_map, ['マスタ名', '台帳名', '変換前'], 0)
                col_freee = find_col(df_map, ['Freee名', 'フリー名', '変換後'], 1)
                
                for _, row in df_map.iterrows():
                    m_name = str(row.get(col_master, "")).strip()
                    f_name = str(row.get(col_freee, "")).strip()
                    if m_name and f_name and m_name != 'nan' and f_name != 'nan':
                        # 比較用に正規化した名前をキー・バリューにする
                        m_clean = normalize_name(m_name)
                        f_clean = normalize_name(f_name)
                        mapping_dict[m_clean] = f_clean
                log_cb(f"名前変換マスタから {len(mapping_dict)} 件の変換ルールを登録しました。")
        except Exception as e:
            log_cb(f"[エラー] 名前変換マスタの読み込みに失敗しました: {e}")
            
    return mapping_dict

# ---------------------------------------------------------
# 進捗・エラーログ用カスタムダイアログ
# ---------------------------------------------------------
class ProgressDialog(tk.Toplevel):
    def __init__(self, parent, title="処理中..."):
        super().__init__(parent)
        self.title(title)
        self.geometry("580x420")
        self.resizable(False, False)
        self.attributes('-topmost', True)
        self.protocol("WM_DELETE_WINDOW", self.on_close)
        
        self.canceled = False
        
        # ステータス表示
        self.lbl_status = tk.Label(self, text="準備中...", font=("Meiryo", 9, "bold"), anchor="w", justify="left")
        self.lbl_status.pack(fill="x", padx=20, pady=(15, 5))
        
        # 進捗バー
        self.progress = ttk.Progressbar(self, orient="horizontal", mode="determinate")
        self.progress.pack(fill="x", padx=20, pady=5)
        
        # 件数表示
        self.lbl_count = tk.Label(self, text="0 / 0 件", font=("Meiryo", 9), anchor="e")
        self.lbl_count.pack(fill="x", padx=20, pady=2)
        
        # 時間情報
        self.lbl_time = tk.Label(self, text="経過時間: 00:00 (残り予想時間: 計算中)", font=("Meiryo", 9), anchor="w")
        self.lbl_time.pack(fill="x", padx=20, pady=2)
        
        # ログ／エラーテキスト表示エリア
        tk.Label(self, text="実行ログ / エラー情報:", font=("Meiryo", 9), anchor="w").pack(fill="x", padx=20, pady=(10, 2))
        
        log_frame = tk.Frame(self)
        log_frame.pack(fill="both", expand=True, padx=20, pady=2)
        
        self.txt_log = tk.Text(log_frame, font=("Consolas", 9), state="disabled", wrap="word")
        self.txt_log.pack(side="left", fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(log_frame, command=self.txt_log.yview)
        scrollbar.pack(side="right", fill="y")
        self.txt_log.config(yscrollcommand=scrollbar.set)
        
        # ボタン領域
        self.btn_frame = tk.Frame(self)
        self.btn_frame.pack(fill="x", padx=20, pady=15)
        
        self.btn_cancel = tk.Button(self.btn_frame, text="処理を中断", command=self.on_cancel, width=15, bg="#ffebee", font=("Meiryo", 9))
        self.btn_cancel.pack(side="right", padx=5)
        
        self.btn_copy = tk.Button(self.btn_frame, text="ログをコピー", command=self.copy_log, width=15, state="disabled", font=("Meiryo", 9))
        self.btn_copy.pack(side="right", padx=5)

    def append_log(self, text):
        self.txt_log.config(state="normal")
        self.txt_log.insert(tk.END, text + "\n")
        self.txt_log.see(tk.END)
        self.txt_log.config(state="disabled")
        self.btn_copy.config(state="normal")

    def on_cancel(self):
        if not self.canceled:
            if messagebox.askyesno("中断確認", "本当に処理を中断しますか？"):
                self.canceled = True
                self.append_log("\n--- ユーザーにより処理が中断されました ---")
                self.btn_cancel.config(state="disabled")

    def on_close(self):
        if self.btn_cancel.cget("text") == "閉じる":
            self.destroy()
        else:
            self.on_cancel()
        
    def copy_log(self):
        self.clipboard_clear()
        self.clipboard_append(self.txt_log.get("1.0", tk.END))
        messagebox.showinfo("完了", "ログをクリップボードにコピーしました。")

def run_with_progress(parent, title, task_func):
    """
    バックグラウンドスレッドで重いタスクを実行し、進捗ダイアログを制御するラッパー。
    """
    dialog = ProgressDialog(parent, title=title)
    q = queue.Queue()
    start_time = time.time()
    
    def worker():
        try:
            def progress_cb(current, total, status_text):
                q.put({'type': 'progress', 'current': current, 'total': total, 'status': status_text})
                if dialog.canceled:
                    raise InterruptedError("ユーザーによる中断要求")
            
            def log_cb(text):
                q.put({'type': 'log', 'text': text})
                
            task_func(progress_cb, log_cb)
            q.put({'type': 'done'})
        except InterruptedError as ie:
            q.put({'type': 'log', 'text': f"処理が中断されました: {ie}"})
            q.put({'type': 'done'})
        except Exception as e:
            tb = traceback.format_exc()
            q.put({'type': 'error', 'text': tb})
            q.put({'type': 'done'})

    # スレッド起動
    t = threading.Thread(target=worker)
    t.daemon = True
    t.start()
    
    def update_gui():
        while not q.empty():
            msg = q.get_nowait()
            if msg['type'] == 'progress':
                c = msg['current']
                tot = msg['total']
                status = msg['status']
                
                dialog.lbl_status.config(text=status)
                if tot > 0:
                    dialog.progress.config(maximum=tot, value=c)
                    dialog.lbl_count.config(text=f"{c} / {tot} 件")
                    
                    elapsed = time.time() - start_time
                    elapsed_str = f"{int(elapsed // 60):02d}:{int(elapsed % 60):02d}"
                    
                    if c > 0:
                        avg_time = elapsed / c
                        remaining = avg_time * (tot - c)
                        rem_str = f"{int(remaining // 60):02d}:{int(remaining % 60):02d}"
                    else:
                        rem_str = "計算中"
                    
                    dialog.lbl_time.config(text=f"経過時間: {elapsed_str} (残り予想時間: {rem_str})")
            
            elif msg['type'] == 'log':
                dialog.append_log(msg['text'])
                
            elif msg['type'] == 'error':
                dialog.append_log("\n========================================\n!!! エラー（バグ）を検出しました !!!\n========================================\n" + msg['text'])
                dialog.lbl_status.config(text="エラー終了")
                messagebox.showerror("エラー検出", "処理中にエラー（バグ）が発生しました。\nログ欄のトレースバックを確認してください。")
                
            elif msg['type'] == 'done':
                dialog.btn_cancel.config(state="normal", text="閉じる")
                dialog.btn_cancel.config(command=dialog.destroy, bg="#e0e0e0")
                if not dialog.canceled and "エラー" not in dialog.lbl_status.cget("text"):
                    dialog.lbl_status.config(text="すべての処理が正常に完了しました")
                    dialog.append_log("\n--- 処理完了 ---")
                dialog.update() # 完了時の強制描画
                return # 監視終了
                
        dialog.update_idletasks() # 強制描画更新
        parent.after(100, update_gui)

    dialog.update() # 初期表示の強制描画
    parent.after(100, update_gui)

# ---------------------------------------------------------
# 【STEP-1】: PC名と氏名の紐づけ（提出用シート作成）
# ---------------------------------------------------------
def step1_create_submission_list(csv_paths, master_path, output_dir, progress_cb, log_cb):
    # 除外マスタのロード
    excludes = load_exclude_list(master_path, log_cb)
    
    log_cb("管理コンソール(台帳)を読み込んでいます...")
    progress_cb(0, len(csv_paths) + 2, "台帳の読み込み中...")
    df_master = load_master_safe(master_path)
    
    dict_terminal = {}
    dict_dept = {}
    for _, row in df_master.iterrows():
        if len(row) > 4:
            comp_name = str(row.iloc[4]).strip().upper()
            if comp_name and comp_name != 'NAN':
                dict_terminal[comp_name] = str(row.iloc[3]) if len(row) > 3 else ""
                dict_dept[comp_name] = str(row.iloc[9]) if len(row) > 9 else ""
    
    log_cb(f"台帳から {len(dict_terminal)} 台の端末情報を取得しました。")
    
    log_cb("SKYSEAログから電源ON/OFF時間を抽出中...")
    results = {}
    for idx, csv_path in enumerate(csv_paths):
        filename = os.path.basename(csv_path)
        progress_cb(idx + 1, len(csv_paths) + 2, f"電源ログを解析中 ({idx + 1}/{len(csv_paths)})")
        log_cb(f"ファイル処理中: {filename}")
        
        df_log = load_csv_safely(csv_path)
        for _, row in df_log.iterrows():
            if len(row) > 11:
                dt_val = str(row.iloc[6]).strip()            # G列: 日時
                comp_name = str(row.iloc[1]).strip().upper() # B列: コンピューター名
                op_type = str(row.iloc[11])                  # L列: 操作種別
                
                # 除外判定 (コンピューター名による除外)
                if comp_name in excludes:
                    continue
                
                if not dt_val or dt_val.lower() == 'nan' or not comp_name or comp_name.lower() == 'nan':
                    continue
                
                try:
                    dt = pd.to_datetime(dt_val)
                    logical_dt = dt - pd.Timedelta(hours=5)
                    yobi = ["月", "火", "水", "木", "金", "土", "日"][logical_dt.weekday()]
                    date_str = logical_dt.strftime('%Y/%m/%d') + f"（{yobi}）"
                    key = (date_str, comp_name)
                    
                    if key not in results:
                        results[key] = [None, None]
                        
                    # PCの起動や操作開始を表す操作種別
                    is_on_event = any(kw in op_type for kw in ["電源ON", "操作開始", "ログオン", "ロック解除", "レジューム"])
                    # PCの終了や操作終了を表す操作種別
                    is_off_event = any(kw in op_type for kw in ["電源OFF", "操作終了", "ログオフ", "ロック", "サスペンド"])
                    
                    if is_on_event:
                        if results[key][0] is None or dt < results[key][0]:
                            results[key][0] = dt
                    if is_off_event:
                        if results[key][1] is None or dt > results[key][1]:
                            results[key][1] = dt
                except:
                    pass
                    
    progress_cb(len(csv_paths) + 1, len(csv_paths) + 2, "出力用CSVデータを構築中...")
    log_cb("CSVデータ変換・除外処理を行っています...")
    
    output_data = []
    excluded_count = 0
    for (date_str, comp_name), (time_on_dt, time_off_dt) in results.items():
        term_name = dict_terminal.get(comp_name, "(未登録)")
        dept_name = dict_dept.get(comp_name, "-")
        
        # 氏名（端末機名）による除外判定
        term_name_clean = re.sub(r'\s+', '', term_name).upper()
        if term_name_clean in excludes:
            excluded_count += 1
            continue
            
        time_on = time_on_dt.strftime('%H:%M:%S') if time_on_dt is not None else ""
        time_off = time_off_dt.strftime('%H:%M:%S') if time_off_dt is not None else ""
        output_data.append([date_str, term_name, comp_name, dept_name, time_on, time_off])
        
    log_cb(f"除外マスタに基づき、合計 {excluded_count} 行の共用・業務用端末ログを除外しました。")
        
    df_out = pd.DataFrame(output_data, columns=["日付", "端末機名", "コンピューター名", "部署名", "電源ON時間", "電源OFF時間"])
    
    # 部署名 -> コンピューター名 -> 端末機名 -> 日付 の優先順位で並び替え
    df_out = df_out.sort_values(by=["部署名", "コンピューター名", "端末機名", "日付"]).reset_index(drop=True)
    
    progress_cb(len(csv_paths) + 2, len(csv_paths) + 2, "CSVファイルに書き出し中...")
    out_filename = f"管理部提出用リスト_{time.strftime('%Y%m%d%H%M%S')}.csv"
    out_filepath = os.path.join(output_dir, out_filename)
    log_cb(f"新規CSV作成中: {out_filepath}")
    
    df_out.to_csv(out_filepath, index=False, encoding='cp932')
    log_cb(f"保存完了: {out_filename}")

# ---------------------------------------------------------
# 【STEP-2】: ONOFF時間とFreee打刻の突合整理
# ---------------------------------------------------------
def step2_create_matching_list(step1_csv_path, freee_csvs, output_dir, progress_cb, log_cb):
    total_steps = len(freee_csvs) + 3
    curr_step = 0
    
    log_cb("ステップ１の出力データ(提出用CSV)を読み込んでいます...")
    progress_cb(curr_step, total_steps, "提出用CSVのロード中...")
    
    df_step1 = load_csv_safely(step1_csv_path)
    df_step1.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_step1.columns]
    curr_step += 1
    log_cb("Freee打刻履歴の抽出中...")
    freee_stamps = {}
    for csv_idx, freee_csv in enumerate(freee_csvs):
        progress_cb(curr_step, total_steps, f"Freeeログを解析中 ({csv_idx + 1}/{len(freee_csvs)})")
        log_cb(f"ファイル処理中: {os.path.basename(freee_csv)}")
        
        df_log = load_csv_safely(freee_csv)
        df_log.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_log.columns]
        
        col_comp = find_col(df_log, ['コンピューター名', 'PC名', 'ホスト名'], 1)
        col_datetime = find_col(df_log, ['日時', '日付時刻', '時刻'], 6)
        col_path = find_col(df_log, ['パス', 'URL', 'ＵＲＬ'], 11)
        col_title = df_log.columns[14] if len(df_log.columns) > 14 else find_col(df_log, ['タイトル'], 14)
        
        df_log = df_log.sort_values(by=[col_comp, col_datetime]).reset_index(drop=True)
        
        for i in range(len(df_log) - 1):
            row_curr = df_log.iloc[i]
            row_next = df_log.iloc[i + 1]
            
            comp_curr = str(row_curr.get(col_comp, '')).strip().upper()
            comp_next = str(row_next.get(col_comp, '')).strip().upper()
            
            if comp_curr != comp_next or comp_curr == '' or comp_curr == 'NAN':
                continue
            
            path_curr = str(row_curr.get(col_path, '')).strip()
            title_curr = str(row_curr.get(col_title, '')).strip()
            path_next = str(row_next.get(col_path, '')).strip()
            title_next = str(row_next.get(col_title, '')).strip()
            
            is_login = path_curr.startswith(FREEE_LOGIN_URL) and FREEE_LOGIN_TITLE in title_curr
            is_stamp = path_next.startswith(FREEE_STAMP_URL) and FREEE_STAMP_TITLE in title_next
            
            if is_login and is_stamp:
                try:
                    dt_stamp = pd.to_datetime(row_next[col_datetime])
                    logical_dt = dt_stamp - pd.Timedelta(hours=5)
                    date_str = logical_dt.strftime('%Y/%m/%d')
                    key = (date_str, comp_curr)
                    
                    if key not in freee_stamps:
                        freee_stamps[key] = []
                    freee_stamps[key].append(dt_stamp)
                except:
                    pass
        curr_step += 1

    progress_cb(curr_step, total_steps, "データの突合・例外パターン判定中...")
    log_cb("提出用データの電源時間とFreee打刻時間を突き合わせています...")
    
    output_rows = []
    
    for _, row in df_step1.iterrows():
        raw_date = get_val(row, '日付', 0)
        clean_date = raw_date.split('（')[0].split('(')[0].strip()
        
        comp_name = get_val(row, 'コンピューター名', 2).strip().upper()
        name = get_val(row, '端末機名', 1)
        dept = get_val(row, '部署名', 3)
        
        on_str = get_val(row, '電源ON時間', 4).strip()
        off_str = get_val(row, '電源OFF時間', 5).strip()
        
        on_dt = None
        off_dt = None
        
        try:
            if on_str and on_str != 'nan':
                on_dt = pd.to_datetime(f"{clean_date} {on_str}")
        except:
            pass
            
        try:
            if off_str and off_str != 'nan':
                off_dt = pd.to_datetime(f"{clean_date} {off_str}")
        except:
            pass
            
        key = (clean_date, comp_name)
        stamps = sorted(freee_stamps.get(key, []))
        
        in_op_time = ""
        out_op_time = ""
        
        if stamps:
            # PC出勤操作時刻の判定
            if on_dt:
                after_on = [s for s in stamps if s >= on_dt]
                in_op_time = after_on[0].strftime('%H:%M:%S') if after_on else stamps[0].strftime('%H:%M:%S')
            else:
                in_op_time = stamps[0].strftime('%H:%M:%S')
            
            # PC退勤操作時刻の判定
            if off_dt:
                before_off = [s for s in stamps if s <= off_dt]
                out_op_time = before_off[-1].strftime('%H:%M:%S') if before_off else stamps[-1].strftime('%H:%M:%S')
            else:
                out_op_time = stamps[-1].strftime('%H:%M:%S')
        else:
            # 例外B: Freeeログなし（スマホ打刻）
            in_op_time = on_str if on_str != 'nan' else ""
            out_op_time = off_str if off_str != 'nan' else ""
            
        output_rows.append([raw_date, name, comp_name, dept, on_str if on_str != 'nan' else "", off_str if off_str != 'nan' else "", in_op_time, out_op_time])
        
    df_out = pd.DataFrame(output_rows, columns=[
        "日付", "氏名", "コンピューター名", "部署名", 
        "PC電源ON時間", "PC電源OFF時間", "PC出勤操作時刻", "PC退勤操作時刻"
    ])
    
    # 部署名 -> コンピューター名 -> 氏名 -> 日付 の優先順位で並び替え
    df_out = df_out.sort_values(by=["部署名", "コンピューター名", "氏名", "日付"]).reset_index(drop=True)
    curr_step += 1
    
    progress_cb(curr_step, total_steps, "Excel突合リスト出力中...")
    out_filename = f"出退勤突合リスト_{time.strftime('%Y%m%d%H%M%S')}.xlsx"
    out_filepath = os.path.join(output_dir, out_filename)
    log_cb(f"新規Excel作成中: {out_filepath}")
    
    df_out.to_excel(out_filepath, index=False, sheet_name="出退勤突合結果")
    
    wb = openpyxl.load_workbook(out_filepath)
    ws = wb["出退勤突合結果"]
    fill_color = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
    for cell in ws[1]:
        cell.fill = fill_color
        
    for col in ws.columns:
        max_length = 0
        column = get_column_letter(col[0].column)
        for cell in col:
            try:
                if len(str(cell.value)) > max_length: max_length = len(str(cell.value))
            except: pass
        ws.column_dimensions[column].width = (max_length + 2)
        
    wb.save(out_filepath)
    log_cb(f"保存完了: {out_filename}")

# ---------------------------------------------------------
# 【STEP-3】: 乖離・詳細調査（Freee実データ突合・判定）
# ---------------------------------------------------------
def step3_analyze_gap(matching_xlsx, freee_work_path, output_dir, progress_cb, log_cb):
    try:
        progress_cb(0, 100, "データの読み込み中...")
        log_cb("突合リストを読み込んでいます...")
        df_match = pd.read_excel(matching_xlsx)
        
        log_cb("日次勤怠一覧データを読み込んでいます...")
        df_work = load_master_safe(freee_work_path)
        
        progress_cb(5, 100, "データのクレンジングと紐付け中...")
        
        # 列名（ヘッダー）のキーワード自動検知（インデックスズレ対策）
        cols_needed = {
            '氏名_work': find_col(df_work, ['氏名', '名前', '従業員'], 1),
            '日付_work': find_col(df_work, ['日付', '年月日', '日'], 5),
            '部門_work': find_col(df_work, ['部門', '部署', '組織'], 3),
            '勤務予定開始時刻': find_col(df_work, ['予定開始', '勤務予定開始', '始業予定'], 8),
            '勤務予定退勤時刻': find_col(df_work, ['予定退勤', '勤務予定退勤', '終業予定'], 9),
            '出勤時刻（実時間）': find_col(df_work, ['出勤時刻', '実出勤', '始業時刻'], 11),
            '退勤時刻（実時間）': find_col(df_work, ['退勤時刻', '実退勤', '終業時刻'], 12),
            '休憩時間（実時間）': find_col(df_work, ['休憩時間', '休憩'], 22),
            '総勤務時間（実時間）': find_col(df_work, ['総勤務時間', '総労働', '実労働'], 23)
        }
        
        log_cb(f"勤怠シートの該当列名を取得しました。")

        # 名前変換マスタのロード
        name_map = load_name_mapping(freee_work_path, log_cb)
        
        # 突合側のキー作成
        df_match['氏名_clean'] = df_match['氏名'].apply(normalize_name)
        df_match['氏名_key'] = df_match['氏名_clean'].apply(lambda x: name_map.get(x, x))
        
        # 日付文字列から (火) などの曜日部分を除去して YYYY/MM/DD 形式にする
        df_match['日付_key'] = df_match['日付'].astype(str).apply(lambda x: x.split('（')[0].split('(')[0].strip())
        df_match['日付_key'] = pd.to_datetime(df_match['日付_key'], errors='coerce').dt.strftime('%Y/%m/%d')
        
        # 勤怠側のキー作成
        df_work['氏名_key'] = df_work[cols_needed['氏名_work']].apply(normalize_name)
        
        # 日次勤怠一覧の日付を YYYY/MM/DD に統一
        df_work['日付_dt'] = pd.to_datetime(df_work[cols_needed['日付_work']], errors='coerce')
        df_work = df_work.dropna(subset=['日付_dt', '氏名_key'])
        df_work['日付_key'] = df_work['日付_dt'].dt.strftime('%Y/%m/%d')
        df_work = df_work.drop(columns=['日付_dt'])
        
        # マージ前にFreee側のデータ行が存在することを示すフラグを追加
        df_work['has_freee_flag'] = True
        
        log_cb(f"マージ前: 突合側 {len(df_match)} 件, Freee側 {len(df_work)} 件")
        
        # Freeeのデータを基準とする結合 (Left Join) - 同名の列が衝突した場合は _work / _match を付与
        df_merged = pd.merge(df_work, df_match, on=['日付_key', '氏名_key'], how='left', suffixes=('_work', '_match'))
        
        log_cb(f"マージ後: {len(df_merged)} 件")
        
        # マージ後の日付・氏名補完処理
        def format_date_with_yobi(date_key):
            try:
                dt = pd.to_datetime(date_key)
                yobi = ["月", "火", "水", "木", "金", "土", "日"][dt.weekday()]
                return dt.strftime('%Y/%m/%d') + f"（{yobi}）"
            except:
                return date_key
                
        df_merged['日付'] = df_merged['日付_key'].apply(format_date_with_yobi)
        
        # 安全に列名を取得する補助関数
        def get_merged_col(df, col_name, suffix='_work'):
            if col_name in df.columns:
                return df[col_name]
            elif f"{col_name}{suffix}" in df.columns:
                return df[f"{col_name}{suffix}"]
            return pd.Series(index=df.index, dtype='object') # 存在しない場合は空のSeries
            
        # 氏名補完: Freee側の登録名(姓名スペースありなど)を最優先し、無ければ突合側を採用
        match_name = get_merged_col(df_merged, '氏名', '_match')
        work_name = get_merged_col(df_merged, cols_needed['氏名_work'], '_work')
        df_merged['氏名'] = work_name.fillna(match_name)
        
        # 部署名・部門補完（Freee側の部署名を優先）
        match_dept = get_merged_col(df_merged, '部署名', '_match')
        work_dept = get_merged_col(df_merged, cols_needed['部門_work'], '_work')
        df_merged['部署名'] = work_dept.fillna(match_dept).fillna('-')
        
        # 列名マッピング
        # Freeeの列をExcelに出力するための名前にコピーする
        df_merged['勤務予定開始時刻'] = get_merged_col(df_merged, cols_needed['勤務予定開始時刻'], '_work')
        df_merged['勤務予定退勤時刻'] = get_merged_col(df_merged, cols_needed['勤務予定退勤時刻'], '_work')
        df_merged['出勤時刻（実時間）'] = get_merged_col(df_merged, cols_needed['出勤時刻（実時間）'], '_work')
        df_merged['退勤時刻（実時間）'] = get_merged_col(df_merged, cols_needed['退勤時刻（実時間）'], '_work')
        df_merged['休憩時間（実時間）'] = get_merged_col(df_merged, cols_needed['休憩時間（実時間）'], '_work')
        df_merged['総勤務時間（実時間）'] = get_merged_col(df_merged, cols_needed['総勤務時間（実時間）'], '_work')
        
        # 新設列の初期化
        df_merged['早出乖離'] = ""
        df_merged['残業乖離'] = ""
        df_merged['乖離判定'] = "問題なし"
        df_merged['調査依頼'] = ""
        
        total_rows = len(df_merged)
        
        # 乖離判定ループ
        for idx, row in df_merged.iterrows():
            if idx % 100 == 0 or idx == total_rows - 1:
                progress_cb(int((idx / total_rows) * 90) + 5, 100, f"乖離データを判定中... ({idx+1}/{total_rows})")
                
            has_freee = row.get('has_freee_flag') == True
            
            # PCログの有無確認
            pc_on_str = str(row.get('PC電源ON時間', '')).strip()
            pc_off_str = str(row.get('PC電源OFF時間', '')).strip()
            has_pc_log = (pc_on_str and pc_on_str.lower() != 'nan') or (pc_off_str and pc_off_str.lower() != 'nan')
            
            # Freeeに情報がない場合
            if not has_freee:
                df_merged.at[idx, '乖離判定'] = "Freeeに情報がありません"
                continue
                
            # 公休日の判定
            def is_empty_time(val):
                if pd.isna(val): return True
                s = str(val).strip()
                return s == "" or s == "-" or s.lower() == "nan"
                
            no_plan = is_empty_time(row['勤務予定開始時刻']) and is_empty_time(row['勤務予定退勤時刻'])
            no_record = is_empty_time(row['出勤時刻（実時間）']) and is_empty_time(row['退勤時刻（実時間）'])
            is_no_work_day_candidate = no_plan and no_record
            
            # 土日判定
            date_key = row['日付_key']
            is_weekend = False
            try:
                dt_val = pd.to_datetime(date_key)
                is_weekend = dt_val.weekday() in [5, 6] # 5:土, 6:日
            except:
                pass
                
            # 予定・実績が空で、かつ土日の場合のみ公休日とする
            is_no_work_day = is_no_work_day_candidate and is_weekend
            
            # 公休日にPCが動いている場合
            if is_no_work_day:
                if has_pc_log:
                    df_merged.at[idx, '乖離判定'] = "公休日PC稼働の疑い"
                else:
                    df_merged.at[idx, '乖離判定'] = "問題なし"
                continue
                
            # 予定・実績が空だが、平日の場合 (出向やパートなど打刻なし運用の可能性)
            if is_no_work_day_candidate and not is_weekend:
                # 平日のPC稼働は問題なしとして扱う
                df_merged.at[idx, '乖離判定'] = "問題なし"
                continue
                
            # PCログがない場合
            if not has_pc_log:
                if not no_record:
                    # Freeeに勤務実績があるにもかかわらず、PC電源ログがない
                    df_merged.at[idx, '乖離判定'] = "PCの電源を切っていないため取得できませんでした"
                else:
                    df_merged.at[idx, '乖離判定'] = "問題なし"
                continue
                
            # 出勤日の乖離計算
            date_key = row['日付_key']
            work_start_dt = parse_time_with_crossover(date_key, row['出勤時刻（実時間）'])
            pc_on_dt = parse_time_with_crossover(date_key, row['PC電源ON時間'])
            
            # 早出乖離の判定
            early_gap_str = ""
            if work_start_dt is not pd.NaT and pc_on_dt is not pd.NaT:
                diff_sec = (work_start_dt - pc_on_dt).total_seconds()
                if diff_sec >= 900:  # 15分以上
                    diff_mins = int(diff_sec / 60)
                    early_gap_str = format_minutes_to_hm(diff_mins)
                    df_merged.at[idx, '早出乖離'] = early_gap_str
                    
            # 残業乖離の判定
            over_gap_str = ""
            # 日またぎ補正用の基準日時 (出勤開始時間をベースにする、なければ電源ON時間)
            base_ref_dt = work_start_dt if work_start_dt is not pd.NaT else pc_on_dt
            
            work_end_dt = parse_time_with_crossover(date_key, row['退勤時刻（実時間）'], base_ref_dt)
            pc_off_dt = parse_time_with_crossover(date_key, row['PC電源OFF時間'], base_ref_dt)
            
            if work_end_dt is not pd.NaT and pc_off_dt is not pd.NaT:
                diff_sec = (pc_off_dt - work_end_dt).total_seconds()
                if diff_sec >= 900:  # 15分以上
                    diff_mins = int(diff_sec / 60)
                    over_gap_str = format_minutes_to_hm(diff_mins)
                    df_merged.at[idx, '残業乖離'] = over_gap_str
                    
            # 乖離メッセージ構築
            if early_gap_str and over_gap_str:
                df_merged.at[idx, '乖離判定'] = f"サービス早出（{early_gap_str}） / サービス残業（{over_gap_str}）"
            elif early_gap_str:
                df_merged.at[idx, '乖離判定'] = f"サービス早出（{early_gap_str}）"
            elif over_gap_str:
                df_merged.at[idx, '乖離判定'] = f"サービス残業（{over_gap_str}）"
            else:
                df_merged.at[idx, '乖離判定'] = "問題なし"

        # カラムの順番を明示的に指定して並び替え (左から 基本情報 -> Freee -> SkySea -> 判定)
        column_order = [
            "日付", "氏名", "部署名", 
            "勤務予定開始時刻", "勤務予定退勤時刻", "出勤時刻（実時間）", "退勤時刻（実時間）", "休憩時間（実時間）", "総勤務時間（実時間）",
            "コンピューター名", "PC電源ON時間", "PC電源OFF時間", "PC出勤操作時刻", "PC退勤操作時刻",
            "早出乖離", "残業乖離", "乖離判定", "調査依頼"
        ]
        df_merged = df_merged[column_order]
        
        # 部署名 -> 氏名 -> 日付 の優先順位で並び替え（コンピューター名を除外して分断を防ぐ）
        df_merged = df_merged.sort_values(by=["部署名", "氏名", "日付"]).reset_index(drop=True)
        
        # ユーザー要望に合わせたヘッダー名の改良（改行コードの追加、PC名への変更）
        rename_dict = {
            "勤務予定開始時刻": "勤務予定\n開始時刻",
            "勤務予定退勤時刻": "勤務予定\n退勤時刻",
            "出勤時刻（実時間）": "出勤時刻\n（実時間）",
            "退勤時刻（実時間）": "退勤時刻\n（実時間）",
            "休憩時間（実時間）": "休憩時間\n（実時間）",
            "総勤務時間（実時間）": "総勤務時間\n（実時間）",
            "コンピューター名": "PC名",
            "PC電源ON時間": "PC電源\nON時間",
            "PC電源OFF時間": "PC電源\nOFF時間",
            "PC出勤操作時刻": "PC出勤\n操作時刻",
            "PC退勤操作時刻": "PC退勤\n操作時刻"
        }
        df_merged = df_merged.rename(columns=rename_dict)
        
        progress_cb(total_rows, total_rows, "Excelファイル作成中...")
        out_filename = f"乖離調査結果_{time.strftime('%Y%m%d%H%M%S')}.xlsx"
        out_filepath = os.path.join(output_dir, out_filename)
        log_cb(f"新規Excel作成中: {out_filepath}")
        
        # Excelの2行目から書き出し
        df_merged.to_excel(out_filepath, index=False, startrow=1, sheet_name="乖離調査結果")
        
        wb = openpyxl.load_workbook(out_filepath)
        ws = wb["乖離調査結果"]
        
        # 1行目のグループヘッダー
        ws['A1'] = "基本情報"
        ws['D1'] = "Freee抽出データ"
        ws['J1'] = "SkySea抽出ログデータ"
        ws['O1'] = "判定"
        
        # セル結合
        ws.merge_cells('A1:C1')
        ws.merge_cells('D1:I1')
        ws.merge_cells('J1:N1')
        ws.merge_cells('O1:R1')
        
        # スタイル設定
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        
        # 境界線
        thin = Side(border_style="thin", color="D9D9D9")
        header_border = Border(left=thin, right=thin, top=thin, bottom=thin)
        
        # カラー (ソフトトーン)
        gray_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        pink_fill = PatternFill(start_color="F2DCDB", end_color="F2DCDB", fill_type="solid")
        blue_fill = PatternFill(start_color="B4C6E7", end_color="B4C6E7", fill_type="solid")
        green_fill = PatternFill(start_color="C6E0B4", end_color="C6E0B4", fill_type="solid")
        
        font_header = Font(name="Meiryo", size=10, bold=True, color="000000")
        align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        # A~C (基本情報): グレー
        for col_idx in range(1, 4):
            for row_idx in [1, 2]:
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.fill = gray_fill
                cell.font = font_header
                cell.alignment = align_center
                cell.border = header_border
                
        # D~I (Freee): ピンク
        for col_idx in range(4, 10):
            for row_idx in [1, 2]:
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.fill = pink_fill
                cell.font = font_header
                cell.alignment = align_center
                cell.border = header_border
                
        # J~N (SkySea): 青
        for col_idx in range(10, 15):
            for row_idx in [1, 2]:
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.fill = blue_fill
                cell.font = font_header
                cell.alignment = align_center
                cell.border = header_border
                
        # O~R (判定/調査依頼): 緑
        for col_idx in range(15, 19):
            for row_idx in [1, 2]:
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.fill = green_fill
                cell.font = font_header
                cell.alignment = align_center
                cell.border = header_border

        # 土日の項目に対する背景色 (土曜は薄い青、日曜は薄い赤)
        sat_fill_row = PatternFill(start_color="EAF2F8", end_color="EAF2F8", fill_type="solid")
        sun_fill_row = PatternFill(start_color="FDEDEC", end_color="FDEDEC", fill_type="solid")
        
        for r_idx in range(3, ws.max_row + 1):
            date_val = str(ws.cell(row=r_idx, column=1).value or '')
            is_sat = '土' in date_val
            is_sun = '日' in date_val
            
            if not is_sat and not is_sun:
                try:
                    clean_date = date_val.split('（')[0].split('(')[0].strip()
                    dt_val = pd.to_datetime(clean_date)
                    is_sat = dt_val.weekday() == 5
                    is_sun = dt_val.weekday() == 6
                except:
                    pass
                    
            if is_sat or is_sun:
                row_fill = sat_fill_row if is_sat else sun_fill_row
                for c_idx in range(1, 19): # A〜R列
                    ws.cell(row=r_idx, column=c_idx).fill = row_fill
        
        # 行高さ設定
        ws.row_dimensions[1].height = 25
        ws.row_dimensions[2].height = 25
        
        # 自動列幅調整
        for col in ws.columns:
            max_length = 0
            column = get_column_letter(col[0].column)
            # 結合セルの影響を避けるため、データ行である3行目以降を元に幅を決定
            for cell in col[2:]:
                try:
                    if cell.value:
                        val_str = str(cell.value)
                        # 全角文字と半角文字の幅差異を考慮して簡易計算
                        byte_len = len(val_str.encode('utf-8'))
                        actual_len = int((byte_len + len(val_str)) / 2)
                        if actual_len > max_length:
                            max_length = actual_len
                except:
                    pass
            ws.column_dimensions[column].width = max(max_length + 4, 10)
            
        wb.save(out_filepath)
        progress_cb(total_rows, total_rows, "完了")
        log_cb(f"保存完了: {out_filename}")
        
    except Exception as e:
        log_cb(f"エラー終了: {e}")
        raise e

# ---------------------------------------------------------
# 【STEP-4】: 残業調査報告書(Word)の自動生成（AI要約）
# ---------------------------------------------------------
def get_audit_target_dates(excel_path, sheet_name='乖離調査結果'):
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        if sheet_name not in wb.sheetnames: return set()
        ws = wb[sheet_name]
        
        name_col_idx = 2
        date_col_idx = 1
        
        for row in ws.iter_rows(min_row=2, max_row=5):
            for cell in row:
                val = str(cell.value).strip() if cell.value else ''
                if '端末' in val or '氏名' in val or '名前' in val: name_col_idx = cell.column
                if '日付' in val: date_col_idx = cell.column
            if name_col_idx and date_col_idx: break
                
        targets = set()
        ac_col_idx = 18  # R列 (調査依頼)
        
        for row_idx in range(3, ws.max_row + 1):
            ac_cell = ws.cell(row=row_idx, column=ac_col_idx)
            val = str(ac_cell.value).strip() if ac_cell.value else ''
            
            if val in ['〇', '○', 'o', 'O', '丸', '対象']:
                name_val = str(ws.cell(row=row_idx, column=name_col_idx).value or '').strip()
                date_val = str(ws.cell(row=row_idx, column=date_col_idx).value or '').strip()
                if name_val and date_val:
                    try:
                        clean_date = date_val.split('（')[0].split('(')[0].strip()
                        dt = pd.to_datetime(clean_date).strftime('%m/%d')
                        name_clean = normalize_name(name_val).replace('社員', '')
                        targets.add((name_clean, dt))
                    except: pass
        return targets
    except Exception as e:
        print(f"監査対象日（R列の〇）の抽出に失敗しました: {e}")
        return set()

def step4_generate_reports(excel_path, csv_paths, output_dir, progress_cb, log_cb):
    log_cb("提出用シート(勤怠マスタ)をロード中...")
    progress_cb(0, 100, "提出用シートのロード中...")
    
    audit_targets = get_audit_target_dates(excel_path)
    if not audit_targets:
        log_cb("[警告] AC列に「〇」（監査対象）が見つかりませんでした。全データを監査対象として処理します。")
    else:
        log_cb(f"提出用シートから監査対象として {len(audit_targets)} 件の「〇」データを検出しました。")
    
    df_master = pd.read_excel(excel_path, sheet_name='乖離調査結果', header=1)
    df_master.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_master.columns]
    
    name_col = find_col(df_master, ['端末', '氏名', '名前'], 1)
    
    log_cb(f"対象者 {len(csv_paths)} 個の操作ログファイルを順次解析します...")
    
    total_files = len(csv_paths)
    for idx_file, csv_path in enumerate(csv_paths):
        filename_csv = os.path.basename(csv_path)
        log_cb(f"----------------------------------------\n[ファイル解析開始 {idx_file+1}/{total_files}]: {filename_csv}")
        
        target_name_guess = None
        for name in df_master[name_col].dropna().unique():
            name_clean = normalize_name(str(name)).replace('社員', '')
            if name_clean and name_clean in normalize_name(filename_csv):
                target_name_guess = str(name)
                break
                
        if not target_name_guess:
            target_name_guess = filename_csv.split('_')[0] if '_' in filename_csv else filename_csv[:2]
            
        df_target_master = df_master[df_master[name_col].str.contains(target_name_guess, na=False)].copy()
        
        if len(df_target_master) == 0:
            log_cb(f"  -> [スキップ] マスタから対象者「{target_name_guess}」を特定できませんでした。")
            continue
            
        date_col = find_col(df_target_master, ['日付'], 0)
        clean_dates = df_target_master[date_col].astype(str).str.split('（').str[0].str.split('(').str[0].str.strip()
        df_target_master['日付_master'] = pd.to_datetime(clean_dates, errors='coerce')
        
        try:
            df_log = load_csv_safely(csv_path)
        except Exception as e:
            log_cb(f"  -> [エラー] CSV読み込み失敗: {e}")
            continue

        df_log.columns = [str(c).strip().replace('\u3000', '').replace(' ', '') for c in df_log.columns]
        col_datetime = find_col(df_log, ['日時'], 6)
        col_path = find_col(df_log, ['パス', 'URL', 'ＵＲＬ'], 11)
        col_title = df_log.columns[14] if len(df_log.columns) > 14 else find_col(df_log, ['タイトル'], 14)
        
        try:
            df_log['日時'] = pd.to_datetime(df_log[col_datetime])
            df_log = df_log.sort_values('日時').reset_index(drop=True)
            df_log['日付_log'] = (df_log['日時'] - pd.Timedelta(hours=5)).dt.normalize()
            
            df_log['次ログ日時'] = df_log['日時'].shift(-1)
            df_log['次ログ日付'] = df_log['日付_log'].shift(-1)
            df_log['操作秒数'] = (df_log['次ログ日時'] - df_log['日時']).dt.total_seconds().fillna(0)
            
            df_log.loc[df_log['日付_log'] != df_log['次ログ日付'], '操作秒数'] = 0
            df_log.loc[df_log['操作秒数'] < 0, '操作秒数'] = 0
            df_log['操作秒数'] = df_log['操作秒数'].apply(lambda x: min(x, 300))
        except Exception as e:
            log_cb(f"  -> [エラー] 日時処理失敗: {e}")
            continue

        daily_log = df_log.groupby('日付_log').agg(PC電源ON=('日時', 'min'), PC電源OFF=('日時', 'max')).reset_index()
        merged_df = pd.merge(df_target_master, daily_log, left_on='日付_master', right_on='日付_log', how='left')
        
        col_start = find_col(df_target_master, ['電源ON時間'], 4)
        col_end = find_col(df_target_master, ['電源OFF時間'], 5)

        valid_rows = []
        for index, row in merged_df.iterrows():
            if pd.isna(row['PC電源ON']): continue
            date_md = row['日付_master'].strftime('%m/%d')
            target_fullname = get_val(row, name_col)
            target_fullname_clean = target_fullname.replace(' ', '').replace('　', '').replace('社員', '')
            if audit_targets and (target_fullname_clean, date_md) not in audit_targets: continue
            valid_rows.append(row)
            
        total_valid = len(valid_rows)
        log_cb(f"  -> 対象者: {target_name_guess} (監査該当数: {total_valid} 日分)")
        
        for idx_row, row in enumerate(valid_rows):
            date_str = row['日付_master'].strftime('%Y/%m/%d')
            target_fullname = get_val(row, name_col)
            
            msg = f"{target_fullname} 氏 [{date_str}] ({idx_row+1}/{total_valid}) AI解析中..."
            progress_val = int(((idx_file + (idx_row / total_valid)) / total_files) * 100)
            progress_cb(progress_val, 100, msg)
            log_cb(f"    [{date_str}] Gemini APIでAI要約を生成中...")
            
            day_logs = df_log[df_log['日付_log'] == row['日付_log']]
            valid_logs = day_logs[day_logs[col_title].notna() & (day_logs[col_title].astype(str).str.strip() != '') & (day_logs[col_title].astype(str) != 'nan')]
            app_summary = valid_logs.groupby(col_title)['操作秒数'].sum().sort_values(ascending=False)
            
            ai_details, ai_summary = generate_summary_with_ai(day_logs, col_title, col_path)
            
            doc = Document()
            doc.add_heading('残業調査報告書', 0)
            yobi = ["月", "火", "水", "木", "金", "土", "日"][row['日付_master'].weekday()]
            
            calc_start = row['PC電源ON']
            calc_end = row['PC電源OFF']
            diff_sec = (calc_end - calc_start).total_seconds()
            est_time_str = f"{calc_start.strftime('%H:%M')} ～ {calc_end.strftime('%H:%M')}（{format_est_time(diff_sec)}）"
            
            doc.add_paragraph(f"対象者：{target_fullname}")
            doc.add_paragraph(f"部署名：{get_val(row, '部署名', 3)}")
            doc.add_paragraph(f"{row['日付_master'].strftime('%Y/%m/%d')}（{yobi}）")
            doc.add_paragraph(f"PCの稼働：{row['PC電源ON'].strftime('%H:%M')} ～ {row['PC電源OFF'].strftime('%H:%M')}（最終）")
            doc.add_paragraph(f"作業していたと推測される時間：{est_time_str}")
            
            doc.add_heading('主な操作内容（タイトル別集計）', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = '操作タイトル'
            table.rows[0].cells[1].text = '合計時間（概算）'
            
            count = 0
            for app_name, total_sec in app_summary.items():
                if count >= 10 or total_sec == 0: break
                row_cells = table.add_row().cells
                title_text = str(app_name)
                if len(title_text) > 40: title_text = title_text[:38] + "..."
                row_cells[0].text = title_text
                
                h_app, rem_app = divmod(int(total_sec), 3600)
                m_app, s_app = divmod(rem_app, 60)
                
                if h_app == 0: row_cells[1].text = f"{m_app}分{s_app}秒"
                elif m_app == 0 and s_app == 0: row_cells[1].text = f"{h_app}時間"
                else: row_cells[1].text = f"{h_app}時間{m_app:02d}分"
                count += 1
                
            doc.add_heading('作業詳細', level=2)
            doc.add_paragraph(ai_details)
            doc.add_heading('サマリー', level=2)
            doc.add_paragraph(ai_summary)
            
            target_name_short = target_fullname.replace('社員', '').replace(' ', '').replace('　', '')
            safe_date = date_str.replace('/', '')
            filename = f"残業調査報告書_{target_name_short}_{safe_date}.docx"
            filepath = os.path.join(output_dir, filename)
            
            try:
                doc.save(filepath)
                log_cb(f"      -> Word保存完了: {filename}")
            except Exception as e:
                log_cb(f"      -> [エラー] Word保存失敗: {e}")
                raise e
            time.sleep(1)

    progress_cb(100, 100, "完了")
    log_cb("すべての報告書の作成が正常に終了しました。")

# ---------------------------------------------------------
# GUI 統合アプリケーション
# ---------------------------------------------------------
class OvertimeSystemApp_V3:
    def __init__(self, root):
        self.root = root
        self.root.title("残業調査 オールインワンシステム Ver.3")
        self.root.geometry("650x380")
        self.root.attributes('-topmost', True)
        
        # タイトル
        tk.Label(root, text="残業調査 監査業務統合システム Ver.3", font=("Meiryo", 14, "bold")).pack(pady=15)
        
        # STEP-1
        btn1 = tk.Button(root, text="ステップ１: PC名と氏名の紐づけ（提出用リスト作成）", command=self.run_step1, width=60, height=2, bg="#e0f7fa", font=("Meiryo", 9, "bold"))
        btn1.pack(pady=5)
        
        # STEP-2
        btn2 = tk.Button(root, text="ステップ２: ONOFF時間とFreee打刻の突合整理", command=self.run_step2, width=60, height=2, bg="#e8f5e9", font=("Meiryo", 9, "bold"))
        btn2.pack(pady=5)
        
        # STEP-3
        btn3 = tk.Button(root, text="ステップ３: 乖離・詳細調査（Freee実データとのマージ）", command=self.run_step3, width=60, height=2, bg="#fffde7", font=("Meiryo", 9, "bold"))
        btn3.pack(pady=5)
        
        # STEP-4
        btn4 = tk.Button(root, text="ステップ４: 残業調査報告書(Word)の自動生成（AI要約）", command=self.run_step4, width=60, height=2, bg="#fce4ec", font=("Meiryo", 9, "bold"))
        btn4.pack(pady=5)

    def run_step1(self):
        messagebox.showinfo("ステップ１", "【1/3】SKYSEAから抽出した「生ログ(CSV/Excel)」を選択してください。(複数選択可)")
        log_csvs = filedialog.askopenfilenames(title="SKYSEA生ログ", filetypes=[("ログファイル", "*.csv *.xlsx *.xls")])
        if not log_csvs: return
        
        messagebox.showinfo("ステップ１", "【2/3】「管理コンソール台帳（CSV/Excel）」を選択してください。")
        master_excel = filedialog.askopenfilename(title="管理コンソール台帳", filetypes=[("Excel/CSV", "*.xlsx *.xls *.csv")])
        if not master_excel: return
        
        messagebox.showinfo("ステップ１", "【3/3】完成した「管理部提出用リスト(CSV)」の保存先フォルダを選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        task = lambda p_cb, l_cb: step1_create_submission_list(log_csvs, master_excel, out_dir, p_cb, l_cb)
        run_with_progress(self.root, "ステップ１: 提出用リスト作成中", task)

    def run_step2(self):
        messagebox.showinfo("ステップ２", "【1/3】ステップ１で作成した「管理部提出用リスト(CSV)」を選択してください。")
        step1_csv = filedialog.askopenfilename(title="管理部提出用リスト(CSV)", filetypes=[("CSV", "*.csv")])
        if not step1_csv: return
        
        messagebox.showinfo("ステップ２", "【2/3】SKYSEAから抽出した「Freeeアクセスログ(CSV)」を選択してください。(複数選択可)")
        freee_csvs = filedialog.askopenfilenames(title="Freeeアクセスログ", filetypes=[("CSVログ", "*.csv")])
        if not freee_csvs: return
        
        messagebox.showinfo("ステップ２", "【3/3】「出退勤突合リスト」の保存先フォルダを選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        task = lambda p_cb, l_cb: step2_create_matching_list(step1_csv, freee_csvs, out_dir, p_cb, l_cb)
        run_with_progress(self.root, "ステップ２: 出退勤突合リスト作成中", task)

    def run_step3(self):
        messagebox.showinfo("ステップ３", "【1/3】ステップ２で出力した「出退勤突合リスト(Excel)」を選択してください。")
        matching_xlsx = filedialog.askopenfilename(title="出退勤突合リスト", filetypes=[("Excel", "*.xlsx")])
        if not matching_xlsx: return
        
        messagebox.showinfo("ステップ３", "【2/3】「Freee勤務実績データ(CSV/Excel)」を選択してください。")
        freee_work = filedialog.askopenfilename(title="Freee勤務実績データ", filetypes=[("CSV/Excel", "*.csv *.xlsx *.xls")])
        if not freee_work: return
        
        messagebox.showinfo("ステップ３", "【3/3】「乖離調査結果」の保存先フォルダを選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        task = lambda p_cb, l_cb: step3_analyze_gap(matching_xlsx, freee_work, out_dir, p_cb, l_cb)
        run_with_progress(self.root, "ステップ３: 乖離・詳細調査中", task)

    def run_step4(self):
        messagebox.showinfo("ステップ４", "【1/3】管理部から返却された「〇」印付きの提出用リスト(Excel)を選択してください。")
        master_excel = filedialog.askopenfilename(title="管理部返却済みリスト(Excel)", filetypes=[("Excel", "*.xlsx *.xls *.xlsm")])
        if not master_excel: return
        
        messagebox.showinfo("ステップ４", "【2/3】対象者の「PC操作ログ(CSV)」を選択してください。(複数選択可)")
        csvs = filedialog.askopenfilenames(title="PC操作ログ(CSV)", filetypes=[("CSV", "*.csv")])
        if not csvs: return
        
        messagebox.showinfo("ステップ４", "【3/3】報告書(Word)を保存する「出力先フォルダ」を選択してください。")
        out_dir = filedialog.askdirectory(title="出力先フォルダ")
        if not out_dir: return
        
        task = lambda p_cb, l_cb: step4_generate_reports(master_excel, csvs, out_dir, p_cb, l_cb)
        run_with_progress(self.root, "ステップ４: Word報告書生成中", task)

def main():
    root = tk.Tk()
    app = OvertimeSystemApp_V3(root)
    root.mainloop()

if __name__ == "__main__":
    main()
