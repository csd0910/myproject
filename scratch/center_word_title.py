from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

files = [
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260404.docx",
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260405.docx",
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260409.docx",
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260410.docx",
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260413.docx",
    r"C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\成川\残業調査報告書_成川真菜_20260415.docx"
]

def center_first_line(file_path):
    if not os.path.exists(file_path):
        print(f"ファイルが見つかりません: {file_path}")
        return

    try:
        doc = Document(file_path)
        if len(doc.paragraphs) > 0:
            # 最初の段落（1行目）を中央揃えに設定
            doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.save(file_path)
            print(f"完了: {os.path.basename(file_path)}")
        else:
            print(f"段落が見つかりません: {os.path.basename(file_path)}")
    except Exception as e:
        print(f"エラー発生 ({os.path.basename(file_path)}): {e}")

if __name__ == "__main__":
    for f in files:
        center_first_line(f)
