from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def update_word():
    path = r'C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\山本光克\残業調査報告書_山本光克_20260413.docx'
    doc = Document()
    
    # タイトル
    title = doc.add_heading('残業調査報告書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本情報
    doc.add_paragraph('対象者：山本光克')
    doc.add_paragraph('2026/04/13（月）')
    doc.add_paragraph('出勤打刻時間：08:27')
    doc.add_paragraph('退勤打刻時間：22:00')
    doc.add_paragraph('PCの稼働：08:27 ～ 00:17（最終）')
    
    # 判定結果
    doc.add_heading('判定結果', level=1)
    doc.add_paragraph('・【残業確認】深夜0:17にPCの「電源OFF」ログを確認しました。')
    doc.add_paragraph('・退勤打刻（22:00）後から0:17までの間、断続的にExcelやPowerPoint、ブラウザ等の操作記録があり、実業務が行われていたものと判断されます。')
    
    # 作業詳細
    doc.add_heading('作業詳細（22:00以降）', level=1)
    details = [
        "・サービス課1課・改善案等資料確認（22:04 ～ 22:21）",
        "・ブラウザによる調査・資料閲覧（22:34 ～ 23:39）",
        "・移動まとめ・コピー作業（23:54 ～ 23:55）",
        "・笹島様向け資料作成 - PowerPoint（00:14 ～ 00:17）",
        "・サイボウズOffice メッセージ詳細確認（00:17）",
        "・深夜 0:17：PC電源OFF（システムログより確認）"
    ]
    for d in details:
        doc.add_paragraph(d)
        
    # サマリー
    doc.add_heading('サマリー', level=1)
    doc.add_paragraph("当日、退勤打刻は22:00に行われているが、その後も翌0:17まで断続的にPC操作が継続されていた。")
    doc.add_paragraph("主な作業内容として、サービス課の資料確認や、深夜には笹島様向けの提出資料（PowerPoint）の編集作業、およびサイボウズOfficeでの連絡確認が行われている。")
    doc.add_paragraph("0:17に明示的な電源OFFログが記録されており、この時刻まで実業務に従事していた実態が認められる。")
    
    doc.save(path)
    print(f"Successfully updated Word report: {path}")

if __name__ == "__main__":
    update_word()
