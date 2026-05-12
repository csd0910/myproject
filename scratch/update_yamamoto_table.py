import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def format_sec(seconds):
    m = int(seconds // 60)
    s = int(seconds % 60)
    if m == 0: return f"{s}秒"
    return f"{m}分{s}秒"

def create_report():
    csv_path = r'C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\山本光克\山本光克_ログ検索結果-20260413.csv'
    df = pd.read_csv(csv_path, encoding='cp932', engine='python')
    df['dt'] = pd.to_datetime(df.iloc[:, 6])
    
    # 04/13 22:00 ～ 04/14 00:17
    target_logs = df[(df['dt'] >= pd.Timestamp('2026-04-13 22:00:00')) & (df['dt'] <= pd.Timestamp('2026-04-14 00:17:00'))].copy()
    target_logs = target_logs.sort_values('dt')
    target_logs['next_dt'] = target_logs['dt'].shift(-1)
    target_logs['diff'] = (target_logs['next_dt'] - target_logs['dt']).dt.total_seconds().fillna(0)
    target_logs['op_sec'] = target_logs['diff'].apply(lambda x: min(x, 300) if x > 0 else 0)
    
    col_title = df.columns[14]
    stats = target_logs[target_logs[col_title].notna()].groupby(col_title).agg(
        total_sec=('op_sec', 'sum'),
        start=('dt', 'min'),
        end=('dt', 'max')
    ).sort_values('total_sec', ascending=False).head(15) # 上位15件
    
    # Word作成
    path = r'C:\Users\フォーレスト026\Desktop\伊藤作業用\残業調査\山本光克\残業調査報告書_山本光克_20260413.docx'
    doc = Document()
    
    title = doc.add_heading('残業調査報告書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('対象者：山本光克')
    doc.add_paragraph('2026/04/13（月）')
    doc.add_paragraph('出勤打刻時間：08:27')
    doc.add_paragraph('退勤打刻時間：22:00')
    doc.add_paragraph('PCの稼働：08:27 ～ 00:17（最終）')
    
    doc.add_heading('判定結果', level=1)
    doc.add_paragraph('・【残業確認】深夜0:17にPCの「電源OFF」ログを確認しました。')
    doc.add_paragraph('・退勤打刻（22:00）後から0:17までの間、断続的にExcelやPowerPoint、ブラウザ等の操作記録があり、実業務が行われていたものと判断されます。')
    
    doc.add_heading('作業詳細', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'タイトル'
    hdr_cells[1].text = '時間帯 / 合計'
    
    for title_text, row in stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(title_text)
        row_cells[1].text = f"{row['start'].strftime('%H:%M')}～{row['end'].strftime('%H:%M')}\n(計 {format_sec(row['total_sec'])})"
    
    doc.add_heading('サマリー', level=1)
    doc.add_paragraph("当日、退勤打刻は22:00に行われているが、その後も翌0:17まで断続的にPC操作が継続されていた。")
    doc.add_paragraph("主な作業内容として、サービス課の資料確認や、深夜には笹島様向けの提出資料（PowerPoint）の編集作業、およびサイボウズOfficeでの連絡確認が行われている。0:17に明示的な電源OFFログが記録されており、この時刻まで実業務に従事していた実態が認められる。")
    
    doc.save(path)
    print(f"Updated: {path}")

if __name__ == "__main__":
    create_report()
